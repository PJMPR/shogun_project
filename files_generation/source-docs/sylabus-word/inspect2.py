# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn

doc = Document('Sylabus3.0_PL.docx')

print(f'Liczba tabel: {len(doc.tables)}')

# Drukuj WSZYSTKIE tabele ze strukturą
for i, table in enumerate(doc.tables):
    print(f'\n=== TABELA {i} ({len(table.rows)} wierszy x {len(table.columns)} kol) ===')
    for j, row in enumerate(table.rows):
        seen = set()
        uniq_cells = []
        for cell in row.cells:
            if id(cell._tc) not in seen:
                seen.add(id(cell._tc))
                uniq_cells.append(cell)
        for k, cell in enumerate(uniq_cells):
            txt = cell.text.strip().replace('\n', ' | ')[:80]
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            span = 1
            if tcPr is not None:
                gs = tcPr.find(qn('w:gridSpan'))
                if gs is not None:
                    span = int(gs.get(qn('w:val'), 1))
            print(f'  R{j}C{k}[s={span}]: {repr(txt)}')

