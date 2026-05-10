# -*- coding: utf-8 -*-
"""
Klasa budująca dokument Word sylabusu na podstawie danych JSON.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import math
import json
from pathlib import Path

from styles import (
    COLOR_HEADER_BG, COLOR_HEADER_DARK, COLOR_WHITE, COLOR_BLACK,
    FONT_NAME, FONT_SIZE_NORMAL, FONT_SIZE_LABEL, FONT_SIZE_TITLE,
    FONT_SIZE_HEADER, MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT, MARGIN_RIGHT
)


def _safe(data, *keys, default=""):
    """Bezpieczne pobieranie zagnieżdżonych kluczy ze słownika."""
    current = data
    for key in keys:
        if isinstance(current, dict):
            current = current.get(key)
        else:
            return default
        if current is None:
            return default
    return current if current is not None else default


def _set_cell_bg(cell, color: str):
    """Ustawia kolor tła komórki tabeli. color = hex string np. 'D9D9D9'"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def _set_table_borders(table):
    """Ustawia obramowanie wszystkich komórek tabeli."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _set_table_width(table, doc):
    """Ustawia szerokość tabeli na 100% strony."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)


def _run_in_cell(cell, text, bold=False, size=None, italic=False):
    """Dodaje run do paragrafu komórki i formatuje go."""
    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size if size else FONT_SIZE_NORMAL
    run.font.bold = bold
    run.font.italic = italic
    return run


def _add_paragraph_to_cell(cell, text, bold=False, size=None, italic=False, first=False):
    """Dodaje nowy paragraf do komórki."""
    if first and cell.paragraphs and cell.paragraphs[0].text == '':
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size if size else FONT_SIZE_NORMAL
    run.font.bold = bold
    run.font.italic = italic
    return p


def _label_cell(cell, text):
    """Formatuje komórkę jako etykietę (szare tło, pogrubiona)."""
    _set_cell_bg(cell, COLOR_HEADER_BG)
    _run_in_cell(cell, text, bold=True, size=FONT_SIZE_LABEL)


def _value_cell(cell, text):
    """Formatuje komórkę jako wartość."""
    _run_in_cell(cell, str(text) if text else "", size=FONT_SIZE_NORMAL)


def _section_header_row(table, text, colspan=None):
    """Dodaje wiersz nagłówka sekcji (szare tło, pogrubiony, scalony)."""
    row = table.add_row()
    cell = row.cells[0]
    if colspan and len(row.cells) > 1:
        cell = row.cells[0]
        for i in range(1, len(row.cells)):
            cell.merge(row.cells[i])
    _set_cell_bg(cell, COLOR_HEADER_BG)
    _run_in_cell(cell, text, bold=True, size=FONT_SIZE_LABEL)
    return row


def _add_section_table(doc, header_text, num_cols=1):
    """Tworzy nową tabelę z nagłówkiem sekcji."""
    table = doc.add_table(rows=0, cols=num_cols)
    _set_table_borders(table)
    _set_table_width(table, doc)
    _section_header_row(table, header_text, colspan=num_cols)
    return table


def _merge_and_set(row, start_col, end_col):
    """Scala komórki od start_col do end_col (włącznie)."""
    cell = row.cells[start_col]
    for i in range(start_col + 1, end_col + 1):
        cell = cell.merge(row.cells[i])
    return cell


# ---------------------------------------------------------------------------
# Ładowanie efektów kierunkowych i budowa odwrotnego mapowania
# ---------------------------------------------------------------------------
_EFEKTY_FILE = Path(__file__).parent.parent / 'public' / 'assets' / 'efekty_ksztalcenia.json'
_efekty_cache: dict | None = None  # {kod_przedmiotu: {kategoria: [K_W01, ...]}}


def _load_efekty_kierunkowe() -> dict:
    """
    Wczytuje efekty_ksztalcenia.json i buduje odwrotne mapowanie:
        { "PAI": {"wiedza": ["K_W09", ...], "umiejetnosci": ["K_U05", ...], "kompetencje_spoleczne": ["K_K01", ...]} }
    """
    global _efekty_cache
    if _efekty_cache is not None:
        return _efekty_cache

    _efekty_cache = {}
    if not _EFEKTY_FILE.exists():
        return _efekty_cache

    try:
        with open(_EFEKTY_FILE, encoding='utf-8') as f:
            data = json.load(f)
    except Exception:
        return _efekty_cache

    efekty = data.get('efekty_ksztalcenia', {})

    # mapowanie nazw kategorii z pliku JSON na klucze używane w sylabus
    kategorie = {
        'wiedza': 'wiedza',
        'umiejetnosci': 'umiejetnosci',
        'kompetencje_spoleczne': 'kompetencje_spoleczne',
    }

    for kat_key, kat_label in kategorie.items():
        for efekt in efekty.get(kat_key, []):
            kod_efektu = efekt.get('kod_efektu', '')
            kody_przedmiotow = efekt.get('kody', [])
            for kod_p in kody_przedmiotow:
                kod_p_upper = str(kod_p).upper()
                if kod_p_upper not in _efekty_cache:
                    _efekty_cache[kod_p_upper] = {'wiedza': [], 'umiejetnosci': [], 'kompetencje_spoleczne': []}
                _efekty_cache[kod_p_upper][kat_label].append(kod_efektu)

    return _efekty_cache


# ---------------------------------------------------------------------------
# Mapowanie kod_przedmiotu → zakres specjalizacyjny
# ---------------------------------------------------------------------------
_ASSETS = Path(__file__).parent.parent / 'public' / 'assets'
_typ_cache: dict | None = None  # {kod_upper: 'obligatoryjny'|'obieralny'|'lektorat'|'specjalizacyjny'}


def _load_typ_przedmiotu() -> dict:
    """
    Buduje mapowanie kod_przedmiotu → zakres specjalizacyjny na podstawie:
      - program.json (obowiązkowe)
      - electives-other.json (obieralne, lektoraty)
      - electives-specializations.json (specjalizacyjne)
    Analogicznie dla katalogu niestacjonarne/.
    Priorytet: specjalizacyjny > lektorat > obieralny > obligatoryjny
    """
    global _typ_cache
    if _typ_cache is not None:
        return _typ_cache

    _typ_cache = {}

    def _process(assets_dir: Path):
        # Grupy lektoratowe — identyfikujemy po id grupy 'LEK*' lub label zawierającym 'lektorat'
        lektorat_ids: set = set()
        other_file = assets_dir / 'electives-other.json'
        if other_file.exists():
            try:
                with open(other_file, encoding='utf-8') as f:
                    other = json.load(f)
                for g in other.get('groups', []):
                    gid = str(g.get('id', '')).upper()
                    label = str(g.get('label', '')).lower()
                    if gid.startswith('LEK') or 'lektorat' in label:
                        lektorat_ids.add(gid)
                    for item in g.get('items', []):
                        kod = str(item.get('code', '')).upper()
                        if gid.startswith('LEK') or 'lektorat' in label:
                            _typ_cache[kod] = 'lektorat'
                        else:
                            # nie nadpisuj jeśli już ustawione jako lektorat
                            if _typ_cache.get(kod) != 'lektorat':
                                _typ_cache[kod] = 'obieralny'
            except Exception:
                pass

        # Specjalizacyjne
        spec_file = assets_dir / 'electives-specializations.json'
        if spec_file.exists():
            try:
                with open(spec_file, encoding='utf-8') as f:
                    spec = json.load(f)
                for sp in spec.get('specializations', []):
                    for item in sp.get('items', []):
                        kod = str(item.get('code', '')).upper()
                        _typ_cache[kod] = 'specjalizacyjny'  # nadpisuje obieralny
            except Exception:
                pass

        # Obowiązkowe z program.json — tylko jeśli kod nie jest jeszcze sklasyfikowany
        prog_file = assets_dir / 'program.json'
        if prog_file.exists():
            try:
                with open(prog_file, encoding='utf-8') as f:
                    prog = json.load(f)
                for sem in prog.get('semesters', []):
                    for s in sem.get('subjects', []):
                        kod = str(s.get('code', '')).upper()
                        if kod not in _typ_cache:
                            _typ_cache[kod] = 'obligatoryjny'
            except Exception:
                pass

    _process(_ASSETS)
    _process(_ASSETS / 'niestacjonarne')
    return _typ_cache


def _get_zakres_specjalizacyjny(kod: str) -> str:
    """Zwraca tekst zakresu specjalizacyjnego dla komórki w tabeli sekcji 7."""
    typ_map = _load_typ_przedmiotu()
    typ = typ_map.get(str(kod).upper(), 'obligatoryjny')
    mapping = {
        'specjalizacyjny': 'specjalizacyjny',
        'lektorat': 'lektorat',
        'obieralny': 'obieralny',
        'obligatoryjny': 'obligatoryjny',
    }
    return mapping.get(typ, 'obligatoryjny')


class SyllabusBuilder:
    """Buduje dokument Word sylabusu na podstawie danych z JSON."""

    def __init__(self, data: dict):
        self.data = data
        self.s = data.get('sylabus', {})
        self.doc = Document()
        self._setup_page()
        # Kody kierunkowych efektów dla tego przedmiotu
        kod = str(self.s.get('kod_przedmiotu', '')).upper()
        efekty_map = _load_efekty_kierunkowe()
        self._kody_kierunkowe = efekty_map.get(kod, {'wiedza': [], 'umiejetnosci': [], 'kompetencje_spoleczne': []})
        # Zakres specjalizacyjny i stopień
        self._zakres = _get_zakres_specjalizacyjny(kod)
        self._stopien = 'I'

    def _setup_page(self):
        """Konfiguruje marginesy strony."""
        for section in self.doc.sections:
            section.top_margin = MARGIN_TOP
            section.bottom_margin = MARGIN_BOTTOM
            section.left_margin = MARGIN_LEFT
            section.right_margin = MARGIN_RIGHT

    def build(self) -> Document:
        """Buduje kompletny dokument sylabusu."""
        self._add_title()
        self._section_1_6()
        self._section_7()
        self._section_8()
        self._section_9()
        self._section_10()
        self._section_11()
        self._section_12()
        self._section_13()
        self._section_14()
        self._section_15()
        self._section_16()
        self._section_17()
        self._section_18_20()
        self._section_21()
        self._section_22()
        self._section_23()
        return self.doc

    # -------------------------------------------------------------------------
    # Nagłówek dokumentu
    # -------------------------------------------------------------------------
    def _add_title(self):
        s = self.s
        uczelnia = _safe(s, 'uczelnia', default='Polsko-Japońska Akademia Technik Komputerowych')
        kierunek = _safe(s, 'kierunek', default='')

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run('SYLABUS')
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_TITLE
        run.font.bold = True

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        run2 = p2.add_run(uczelnia)
        run2.font.name = FONT_NAME
        run2.font.size = Pt(10)
        run2.font.bold = True

        if kierunek:
            p3 = self.doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after = Pt(4)
            run3 = p3.add_run(f'Kierunek: {kierunek}')
            run3.font.name = FONT_NAME
            run3.font.size = Pt(10)

    # -------------------------------------------------------------------------
    # Sekcja 1–6
    # -------------------------------------------------------------------------
    def _section_1_6(self):
        s = self.s
        table = self.doc.add_table(rows=0, cols=4)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        # Wiersz 1: etykiety
        row = table.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        _set_cell_bg(c0, COLOR_HEADER_BG)
        _run_in_cell(c0, '1.  Nazwa polska przedmiotu:', bold=True, size=FONT_SIZE_LABEL)

        c2 = row.cells[2].merge(row.cells[3])
        _set_cell_bg(c2, COLOR_HEADER_BG)
        _run_in_cell(c2, '2.  Wersja z dnia:', bold=True, size=FONT_SIZE_LABEL)

        # Wiersz 1: wartości
        row = table.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        _value_cell(c0, _safe(s, 'nazwa_przedmiotu'))

        c2 = row.cells[2].merge(row.cells[3])
        _value_cell(c2, _safe(s, 'wersja_z_dnia'))

        # Wiersz 2: etykiety
        row = table.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        _set_cell_bg(c0, COLOR_HEADER_BG)
        _run_in_cell(c0, '3.  Nazwa angielska przedmiotu:', bold=True, size=FONT_SIZE_LABEL)

        c2 = row.cells[2].merge(row.cells[3])
        _set_cell_bg(c2, COLOR_HEADER_BG)
        _run_in_cell(c2, '4.  Kod przedmiotu:', bold=True, size=FONT_SIZE_LABEL)

        # Wiersz 2: wartości
        row = table.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        _value_cell(c0, '')  # brak pola w JSON

        c2 = row.cells[2].merge(row.cells[3])
        _value_cell(c2, _safe(s, 'kod_przedmiotu'))

        # Wiersz 3: etykiety
        row = table.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        _set_cell_bg(c0, COLOR_HEADER_BG)
        _run_in_cell(c0, '5.  Nazwa jednostki prowadzącej przedmiot:', bold=True, size=FONT_SIZE_LABEL)

        c2 = row.cells[2].merge(row.cells[3])
        _set_cell_bg(c2, COLOR_HEADER_BG)
        _run_in_cell(c2, '6.  Liczba punktów ECTS:', bold=True, size=FONT_SIZE_LABEL)

        # Wiersz 3: wartości
        row = table.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        _value_cell(c0, _safe(s, 'jednostka'))

        c2 = row.cells[2].merge(row.cells[3])
        _value_cell(c2, str(_safe(s, 'ects')))

    # -------------------------------------------------------------------------
    # Sekcja 7 — Studia
    # -------------------------------------------------------------------------
    def _section_7(self):
        s = self.s
        table = self.doc.add_table(rows=0, cols=6)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        # Nagłówek scalony
        row = table.add_row()
        header_cell = row.cells[0]
        for i in range(1, 6):
            header_cell = header_cell.merge(row.cells[i])
        _set_cell_bg(header_cell, COLOR_HEADER_BG)
        _run_in_cell(header_cell, '7.  Studia', bold=True, size=FONT_SIZE_LABEL)

        # Wiersz nagłówkowy kolumn
        row = table.add_row()
        headers = ['kierunek', 'stopień', 'forma studiów', 'zakres specjalizacyjny', 'profil', 'semestr nauczania']
        for i, h in enumerate(headers):
            _set_cell_bg(row.cells[i], COLOR_HEADER_BG)
            _run_in_cell(row.cells[i], h, bold=True, size=Pt(8))

        # Wiersz danych
        tryb = _safe(s, 'tryb_studiow', default='')
        semestr = _safe(s, 'semestr_studiow', default='')
        profil = _safe(s, 'profil', default='')
        kierunek = _safe(s, 'kierunek', default='')
        values = [kierunek, self._stopien, tryb, self._zakres, profil, str(semestr)]
        row = table.add_row()
        for i, v in enumerate(values):
            _value_cell(row.cells[i], v)

    # -------------------------------------------------------------------------
    # Sekcja 8 — Osoba odpowiedzialna
    # -------------------------------------------------------------------------
    def _section_8(self):
        s = self.s
        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '8.  Osoba odpowiedzialna za przedmiot:', bold=True, size=FONT_SIZE_LABEL)

        row = table.add_row()
        _value_cell(row.cells[0], _safe(s, 'odpowiedzialny_za_przedmiot'))

    # -------------------------------------------------------------------------
    # Sekcja 9 — Formy zajęć  (struktura 1:1 z szablonem Sylabus3.0_PL.docx)
    # Szablon: tabela 10-kolumnowa
    #   R0  [s=10]        nagłówek sekcji
    #   R1  [s=10]        "Wymiar godzinowy..."
    #   R2  [s=2|s=3|s=3|s=2]  nagłówki kolumn godzin
    #   R3  [s=2|s=3|s=3|s=2]  wartości godzin
    #   R4  [s=10]        "Bilans ECTS"
    #   R5  [s=5|s=3|s=2] nagłówki bilansu
    #   R6  [s=5|s=3|s=2] zajęcia z udziałem nauczyciela
    #   R7  [s=5|s=3|s=2] samodzielna praca studenta (liczby)
    #   R8  [s=5|s=5]     samodzielna praca studenta (opis w punktach)
    # -------------------------------------------------------------------------
    def _section_9(self):
        s = self.s
        forma = s.get('forma_i_liczba_godzin_zajec', {})
        godziny = s.get('godziny', {})
        ects_total = _safe(s, 'ects', default=0)

        wyklady = forma.get('wyklady') or 0
        lab = forma.get('laboratorium_projekt') or 0
        cwiczenia = forma.get('cwiczenia_lektorat_seminarium') or 0
        total_cwiczenia = (lab or 0) + (cwiczenia or 0)

        z_udzialem = godziny.get('z_udzialem_prowadzacego_h') or 0
        praca_wlasna_h = godziny.get('praca_wlasna_studenta_h') or 0
        calkowita = godziny.get('calkowita_liczba_godzin_h') or 0
        praca_opis = godziny.get('praca_wlasna_studenta', '') or ''

        ects_z_udzialem = round(z_udzialem / 25, 1) if z_udzialem else 0
        ects_praca = round(praca_wlasna_h / 25, 1) if praca_wlasna_h else 0

        # Jedna tabela 10-kolumnowa (jak w szablonie)
        table = self.doc.add_table(rows=0, cols=10)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        # R0: nagłówek sekcji [s=10]
        row = table.add_row()
        hdr = row.cells[0]
        for i in range(1, 10):
            hdr = hdr.merge(row.cells[i])
        _set_cell_bg(hdr, COLOR_HEADER_BG)
        _run_in_cell(hdr, '9.  Formy zajęć, sposób ich realizacji i przypisana im liczba godzin:', bold=True, size=FONT_SIZE_LABEL)

        # R1: "Wymiar godzinowy..." [s=10]
        row = table.add_row()
        sub = row.cells[0]
        for i in range(1, 10):
            sub = sub.merge(row.cells[i])
        _set_cell_bg(sub, COLOR_HEADER_BG)
        _run_in_cell(sub, 'Wymiar godzinowy (liczba godzin w semestrze)', bold=False, size=Pt(8))

        # R2: nagłówki kolumn godzin [s=2 | s=3 | s=3 | s=2]
        row = table.add_row()
        cells = row.cells
        c0 = cells[0].merge(cells[1])
        c1 = cells[2].merge(cells[3]).merge(cells[4])
        c2 = cells[5].merge(cells[6]).merge(cells[7])
        c3 = cells[8].merge(cells[9])
        for cell, label in [
            (c0, 'wykład'),
            (c1, 'ćwiczenia/laboratorium/pracownia'),
            (c2, 'wykład internetowo'),
            (c3, 'ćw./lab./prac. internetowo'),
        ]:
            _set_cell_bg(cell, COLOR_HEADER_BG)
            _run_in_cell(cell, label, bold=True, size=Pt(8))

        # R3: wartości godzin [s=2 | s=3 | s=3 | s=2]
        row = table.add_row()
        cells = row.cells
        v0 = cells[0].merge(cells[1])
        v1 = cells[2].merge(cells[3]).merge(cells[4])
        v2 = cells[5].merge(cells[6]).merge(cells[7])
        v3 = cells[8].merge(cells[9])
        _value_cell(v0, str(wyklady) if wyklady else '—')
        _value_cell(v1, str(total_cwiczenia) if total_cwiczenia else '—')
        _value_cell(v2, '—')
        _value_cell(v3, '—')

        # R4: "Bilans ECTS" [s=10]
        row = table.add_row()
        cells = row.cells
        bilans_hdr = cells[0]
        for i in range(1, 10):
            bilans_hdr = bilans_hdr.merge(cells[i])
        _set_cell_bg(bilans_hdr, COLOR_HEADER_BG)
        _run_in_cell(bilans_hdr, 'Bilans ECTS', bold=True, size=FONT_SIZE_LABEL)

        # R5: nagłówki bilansu [s=5 | s=3 | s=2]
        row = table.add_row()
        cells = row.cells
        b0 = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3]).merge(cells[4])
        b1 = cells[5].merge(cells[6]).merge(cells[7])
        b2 = cells[8].merge(cells[9])
        for cell, label in [(b0, 'Rodzaj nakładu pracy:'), (b1, 'Liczba godzin'), (b2, 'ECTS')]:
            _set_cell_bg(cell, COLOR_HEADER_BG)
            _run_in_cell(cell, label, bold=True, size=Pt(8))

        # R6: zajęcia z udziałem nauczyciela [s=5 | s=3 | s=2]
        row = table.add_row()
        cells = row.cells
        r0 = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3]).merge(cells[4])
        r1 = cells[5].merge(cells[6]).merge(cells[7])
        r2 = cells[8].merge(cells[9])
        _value_cell(r0, 'Zajęcia z bezpośrednim udziałem nauczyciela')
        _value_cell(r1, str(z_udzialem) if z_udzialem else '')
        _value_cell(r2, str(ects_z_udzialem) if ects_z_udzialem else '')

        # R7: samodzielna praca studenta – liczby [s=5 | s=3 | s=2]
        row = table.add_row()
        cells = row.cells
        s0 = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3]).merge(cells[4])
        s1 = cells[5].merge(cells[6]).merge(cells[7])
        s2 = cells[8].merge(cells[9])
        _value_cell(s0, 'Samodzielna praca studenta')
        _value_cell(s1, str(praca_wlasna_h) if praca_wlasna_h else '')
        _value_cell(s2, str(ects_praca) if ects_praca else '')

        # R8: opis pracy własnej studenta [s=5 | s=5]
        row = table.add_row()
        cells = row.cells
        d0 = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3]).merge(cells[4])
        d1 = cells[5].merge(cells[6]).merge(cells[7]).merge(cells[8]).merge(cells[9])
        _set_cell_bg(d0, COLOR_HEADER_BG)
        _run_in_cell(d0, 'Opis samodzielnej pracy studenta', bold=True, size=Pt(8))
        _value_cell(d1, praca_opis if praca_opis else '')

    # -------------------------------------------------------------------------
    # Sekcja 10 — Język wykładowy
    # -------------------------------------------------------------------------
    def _section_10(self):
        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '10.  Język wykładowy:', bold=True, size=FONT_SIZE_LABEL)

        row = table.add_row()
        _value_cell(row.cells[0], 'polski')

    # -------------------------------------------------------------------------
    # Sekcja 11 — Metody dydaktyczne
    # -------------------------------------------------------------------------
    def _section_11(self):
        s = self.s
        metody = s.get('metody_dydaktyczne', {})

        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '11.  Metody dydaktyczne:', bold=True, size=FONT_SIZE_LABEL)

        row = table.add_row()
        cell = row.cells[0]
        all_methods = []

        _metody_labels = {
            'wyklad': 'Wykład',
            'wykład': 'Wykład',
            'cwiczenia_laboratorium': 'Ćwiczenia / Laboratorium',
            'cwiczenia': 'Ćwiczenia',
            'laboratorium': 'Laboratorium',
        }

        if isinstance(metody, dict):
            for key, methods in metody.items():
                if methods and isinstance(methods, list) and len(methods) > 0:
                    label = _metody_labels.get(key.lower(), key.replace('_', ' ').capitalize())
                    all_methods.append(('label', f'{label}:'))
                    for m in methods:
                        all_methods.append(('item', f'• {m}'))
        elif isinstance(metody, list):
            for m in metody:
                all_methods.append(('item', f'• {m}'))

        first = True
        for kind, text in all_methods:
            if first and cell.paragraphs and cell.paragraphs[0].text == '':
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_NORMAL
            run.font.bold = (kind == 'label')

    # -------------------------------------------------------------------------
    # Sekcja 12 — Zaliczenie i kryteria oceny
    # -------------------------------------------------------------------------
    def _section_12(self):
        s = self.s
        zaliczenie = s.get('zaliczenie', {})
        kryteria = s.get('kryteria_oceny', {})

        # Obsługa nowego formatu {wyklad: [], cwiczenia_laboratorium: []}
        # oraz starego formatu (lista lub string)
        if isinstance(kryteria, dict):
            kryteria_wyklad = kryteria.get('wyklad', []) or []
            kryteria_cw = kryteria.get('cwiczenia_laboratorium', []) or []
        elif isinstance(kryteria, list):
            kryteria_wyklad = kryteria
            kryteria_cw = kryteria
        else:
            kryteria_wyklad = []
            kryteria_cw = []

        table = self.doc.add_table(rows=0, cols=2)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        # Nagłówek scalony
        row = table.add_row()
        header_cell = row.cells[0].merge(row.cells[1])
        _set_cell_bg(header_cell, COLOR_HEADER_BG)
        _run_in_cell(header_cell, '12.  Forma i sposób zaliczenia oraz podstawowe kryteria oceny lub wymagania egzaminacyjne:', bold=True, size=FONT_SIZE_LABEL)

        # Nagłówki kolumn
        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], 'wykład', bold=True, size=Pt(8))
        _set_cell_bg(row.cells[1], COLOR_HEADER_BG)
        _run_in_cell(row.cells[1], 'ćwiczenia/laboratorium/pracownia', bold=True, size=Pt(8))

        # Formy zaliczenia - odczytujemy z zaliczenie{} dopasowując do klucza
        forms = ['Nieoceniany', 'Zaliczenie bez oceny', 'Zaliczenie z oceną', 'Egzamin']
        wyklad_sposob = _safe(zaliczenie, 'Wykład', 'sposob')
        # Szukamy klucza laboratorium lub ćwiczenia
        lab_sposob = (
            _safe(zaliczenie, 'Laboratorium', 'sposob') or
            _safe(zaliczenie, 'Ćwiczenia', 'sposob') or
            _safe(zaliczenie, 'Cwiczenia', 'sposob') or
            _safe(zaliczenie, 'Lektorat', 'sposob') or
            _safe(zaliczenie, 'Seminarium', 'sposob') or ''
        )
        # Jeśli brak klucza Wykład, sprawdź inne warianty
        if not wyklad_sposob:
            wyklad_sposob = next(
                (v.get('sposob', '') for k, v in zaliczenie.items()
                 if 'wyk' in k.lower() and isinstance(v, dict)), ''
            )
        if not lab_sposob:
            lab_sposob = next(
                (v.get('sposob', '') for k, v in zaliczenie.items()
                 if 'wyk' not in k.lower() and isinstance(v, dict)), ''
            )

        row = table.add_row()
        for col_idx, sposob in enumerate([wyklad_sposob, lab_sposob]):
            cell = row.cells[col_idx]
            first = True
            for f in forms:
                if first and cell.paragraphs and cell.paragraphs[0].text == '':
                    p = cell.paragraphs[0]
                    first = False
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                mark = '☑' if f == sposob else '☐'
                run = p.add_run(f'{mark} {f}')
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL

        # Etykiety kryteriów
        row = table.add_row()
        for col_idx in range(2):
            _set_cell_bg(row.cells[col_idx], COLOR_HEADER_BG)
            _run_in_cell(row.cells[col_idx], 'kryteria oceny:', bold=True, size=Pt(8))

        # Wartości kryteriów — osobno dla wykładu i ćwiczeń/lab
        row = table.add_row()
        for col_idx, kryt_list in enumerate([kryteria_wyklad, kryteria_cw]):
            cell = row.cells[col_idx]
            if kryt_list:
                first = True
                for i, k in enumerate(kryt_list):
                    if first and cell.paragraphs and cell.paragraphs[0].text == '':
                        p = cell.paragraphs[0]
                        first = False
                    else:
                        p = cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
                    run = p.add_run(f'{i+1}. {k}')
                    run.font.name = FONT_NAME
                    run.font.size = FONT_SIZE_NORMAL
            else:
                _value_cell(cell, '')

    # -------------------------------------------------------------------------
    # Sekcja 13 — Przedmioty wprowadzające
    # -------------------------------------------------------------------------
    def _section_13(self):
        s = self.s
        przedmioty = s.get('przedmioty_wprowadzajace', [])

        table = self.doc.add_table(rows=0, cols=2)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        # Nagłówek scalony
        row = table.add_row()
        header_cell = row.cells[0].merge(row.cells[1])
        _set_cell_bg(header_cell, COLOR_HEADER_BG)
        _run_in_cell(header_cell, '13.  Kompetencje niezbędne dla realizacji przedmiotu i przedmioty poprzedzające, z których one pochodzą:', bold=True, size=FONT_SIZE_LABEL)

        # Nagłówki kolumn
        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], 'Nazwa przedmiotu poprzedzającego lub realizowanego równolegle', bold=True, size=Pt(8))
        _set_cell_bg(row.cells[1], COLOR_HEADER_BG)
        _run_in_cell(row.cells[1], 'Wymagane kompetencje niezbędne w realizacji danego przedmiotu', bold=True, size=Pt(8))

        if przedmioty:
            for p_item in przedmioty:
                row = table.add_row()
                _value_cell(row.cells[0], _safe(p_item, 'nazwa'))
                _value_cell(row.cells[1], _safe(p_item, 'wymagania'))
        else:
            row = table.add_row()
            _value_cell(row.cells[0], '')
            _value_cell(row.cells[1], '')

    # -------------------------------------------------------------------------
    # Sekcja 14 — Cel dydaktyczny
    # -------------------------------------------------------------------------
    def _section_14(self):
        s = self.s
        cel_pl = _safe(s, 'cel_dydaktyczny')
        cel_eng = _safe(s, 'cel_dydaktyczny_eng')

        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '14.  Cele przedmiotu (abstrakt w języku polskim i angielskim)', bold=True, size=FONT_SIZE_LABEL)

        row = table.add_row()
        cell = row.cells[0]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run('Abstrakt po polsku:')
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_NORMAL
        run.font.bold = True
        if cel_pl:
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_before = Pt(1)
            p2.paragraph_format.space_after = Pt(4)
            run2 = p2.add_run(cel_pl)
            run2.font.name = FONT_NAME
            run2.font.size = FONT_SIZE_NORMAL

        p3 = cell.add_paragraph()
        p3.paragraph_format.space_before = Pt(4)
        p3.paragraph_format.space_after = Pt(1)
        run3 = p3.add_run('Abstrakt po angielsku:')
        run3.font.name = FONT_NAME
        run3.font.size = FONT_SIZE_NORMAL
        run3.font.bold = True
        if cel_eng:
            p4 = cell.add_paragraph()
            p4.paragraph_format.space_before = Pt(1)
            p4.paragraph_format.space_after = Pt(1)
            run4 = p4.add_run(cel_eng)
            run4.font.name = FONT_NAME
            run4.font.size = FONT_SIZE_NORMAL
            run4.font.italic = True

    # -------------------------------------------------------------------------
    # Sekcja 15 — Treści programowe
    # -------------------------------------------------------------------------
    def _section_15(self):
        s = self.s
        tresci = s.get('tresci_programowe', [])

        # Sprawdź czy są jakiekolwiek dane dla ćwiczeń/laboratorium
        has_cw = any(
            str(t.get('cwiczenia') or '').strip()
            for t in tresci if isinstance(t, dict)
        )

        if has_cw:
            # --- Tryb 3-kolumnowy (wykład + ćwiczenia) ---
            table = self.doc.add_table(rows=0, cols=3)
            _set_table_borders(table)
            _set_table_width(table, self.doc)

            def _set_col_width(table, col_widths_cm):
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn as _qn
                tbl = table._tbl
                tblPr = tbl.find(_qn('w:tblPr'))
                existing = tbl.find(_qn('w:tblGrid'))
                if existing is not None:
                    tbl.remove(existing)
                tblGrid = OxmlElement('w:tblGrid')
                for w_cm in col_widths_cm:
                    gridCol = OxmlElement('w:gridCol')
                    gridCol.set(_qn('w:w'), str(int(w_cm * 567)))
                    tblGrid.append(gridCol)
                tbl.insert(list(tbl).index(tblPr) + 1 if tblPr is not None else 0, tblGrid)

            _set_col_width(table, [1.2, 8.0, 7.5])

            row = table.add_row()
            header_cell = row.cells[0]
            for i in range(1, 3):
                header_cell = header_cell.merge(row.cells[i])
            _set_cell_bg(header_cell, COLOR_HEADER_BG)
            _run_in_cell(header_cell, '15.  Treści programowe poszczególnych zajęć:', bold=True, size=FONT_SIZE_LABEL)

            row = table.add_row()
            col_headers = ['Nr zajęć', 'Wykład', 'Ćwiczenia / Laboratorium / Pracownia']
            for i, h in enumerate(col_headers):
                _set_cell_bg(row.cells[i], COLOR_HEADER_BG)
                _run_in_cell(row.cells[i], h, bold=True, size=Pt(8))

            if tresci:
                for i, tresc in enumerate(tresci):
                    row = table.add_row()
                    if isinstance(tresc, dict):
                        nr  = str(tresc.get('nr_zajec', i + 1)) + '.'
                        wyk = str(tresc.get('wyklad') or '')
                        cw  = str(tresc.get('cwiczenia') or '')
                    else:
                        nr = f'{i+1}.'; wyk = str(tresc); cw = ''
                    _value_cell(row.cells[0], nr)
                    _value_cell(row.cells[1], wyk)
                    _value_cell(row.cells[2], cw)
            else:
                for i in range(1, 9):
                    row = table.add_row()
                    _value_cell(row.cells[0], f'{i}.')
                    _value_cell(row.cells[1], '')
                    _value_cell(row.cells[2], '')

        else:
            # --- Tryb 2-kolumnowy (tylko treści programowe) ---
            table = self.doc.add_table(rows=0, cols=2)
            _set_table_borders(table)
            _set_table_width(table, self.doc)

            def _set_col_width2(table, col_widths_cm):
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn as _qn
                tbl = table._tbl
                tblPr = tbl.find(_qn('w:tblPr'))
                existing = tbl.find(_qn('w:tblGrid'))
                if existing is not None:
                    tbl.remove(existing)
                tblGrid = OxmlElement('w:tblGrid')
                for w_cm in col_widths_cm:
                    gridCol = OxmlElement('w:gridCol')
                    gridCol.set(_qn('w:w'), str(int(w_cm * 567)))
                    tblGrid.append(gridCol)
                tbl.insert(list(tbl).index(tblPr) + 1 if tblPr is not None else 0, tblGrid)

            _set_col_width2(table, [1.2, 15.5])

            row = table.add_row()
            header_cell = row.cells[0].merge(row.cells[1])
            _set_cell_bg(header_cell, COLOR_HEADER_BG)
            _run_in_cell(header_cell, '15.  Treści programowe poszczególnych zajęć:', bold=True, size=FONT_SIZE_LABEL)

            row = table.add_row()
            _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
            _set_cell_bg(row.cells[1], COLOR_HEADER_BG)
            _run_in_cell(row.cells[0], 'Nr zajęć', bold=True, size=Pt(8))
            _run_in_cell(row.cells[1], 'Treści programowe', bold=True, size=Pt(8))

            if tresci:
                for i, tresc in enumerate(tresci):
                    row = table.add_row()
                    if isinstance(tresc, dict):
                        nr  = str(tresc.get('nr_zajec', i + 1)) + '.'
                        wyk = str(tresc.get('wyklad') or '')
                    else:
                        nr = f'{i+1}.'; wyk = str(tresc)
                    _value_cell(row.cells[0], nr)
                    _value_cell(row.cells[1], wyk)
            else:
                for i in range(1, 9):
                    row = table.add_row()
                    _value_cell(row.cells[0], f'{i}.')
                    _value_cell(row.cells[1], '')

    # -------------------------------------------------------------------------
    # Sekcja 16 — Elektroniczne materiały dydaktyczne
    # -------------------------------------------------------------------------
    def _section_16(self):
        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '16.  Elektroniczne materiały dydaktyczne', bold=True, size=FONT_SIZE_LABEL)

        row = table.add_row()
        _value_cell(row.cells[0], '')

    # -------------------------------------------------------------------------
    # Sekcja 17 — Literatura
    # -------------------------------------------------------------------------
    def _section_17(self):
        s = self.s
        literatura = s.get('literatura', {})
        podstawowa = _safe(literatura, 'podstawowa', 'pozycje', default=[])
        uzupelniajaca = _safe(literatura, 'uzupelniajaca', 'pozycje', default=[])
        dokumentacja = literatura.get('dokumentacja_internetowa', {})

        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '17.  Wykaz literatury', bold=True, size=FONT_SIZE_LABEL)

        # Literatura wymagana
        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], 'Literatura wymagana do ostatecznego zaliczenia zajęć, wykorzystywana podczas zajęć oraz studiowana samodzielnie przez studenta:', bold=False, size=Pt(8))

        row = table.add_row()
        cell = row.cells[0]
        if podstawowa:
            first = True
            for i, poz in enumerate(podstawowa):
                if first and cell.paragraphs and cell.paragraphs[0].text == '':
                    p = cell.paragraphs[0]
                    first = False
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(f'{i+1}. {poz}')
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL
        else:
            _value_cell(cell, '')

        # Literatura uzupełniająca
        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], 'Literatura uzupełniająca:', bold=False, size=Pt(8))

        row = table.add_row()
        cell = row.cells[0]
        if uzupelniajaca:
            first = True
            for i, poz in enumerate(uzupelniajaca):
                if first and cell.paragraphs and cell.paragraphs[0].text == '':
                    p = cell.paragraphs[0]
                    first = False
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(f'{i+1}. {poz}')
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL
        else:
            _value_cell(cell, '')

        # Dokumentacja internetowa
        if dokumentacja:
            row = table.add_row()
            _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
            _run_in_cell(row.cells[0], 'Dokumentacja internetowa:', bold=False, size=Pt(8))

            row = table.add_row()
            cell = row.cells[0]
            first = True
            for i, (k, v) in enumerate(dokumentacja.items()):
                tekst = v if v else k
                if first and cell.paragraphs and cell.paragraphs[0].text == '':
                    p = cell.paragraphs[0]
                    first = False
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(f'• {tekst}')
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL

    # -------------------------------------------------------------------------
    # Sekcje 18–20 — Efekty kształcenia
    # -------------------------------------------------------------------------
    def _section_18_20(self):
        s = self.s
        efekty = s.get('efekty_ksztalcenia', {})

        sections = [
            ('18.  Wiedza nabywana / dostarczana uczestnikom w trakcie realizacji przedmiotu', 'wiedza'),
            ('19.  Umiejętności nabywane podczas realizacji przedmiotu', 'umiejetnosci'),
            ('20.  Kompetencje społeczne (postawy) kształtowane podczas realizacji przedmiotu', 'kompetencje_spoleczne'),
        ]

        for section_title, key in sections:
            items = efekty.get(key, []) or []

            table = self.doc.add_table(rows=0, cols=3)
            _set_table_borders(table)
            _set_table_width(table, self.doc)

            # Nagłówek scalony
            row = table.add_row()
            header_cell = row.cells[0]
            for i in range(1, 3):
                header_cell = header_cell.merge(row.cells[i])
            _set_cell_bg(header_cell, COLOR_HEADER_BG)
            _run_in_cell(header_cell, section_title, bold=True, size=FONT_SIZE_LABEL)

            # Nagłówki kolumn
            row = table.add_row()
            col_headers = [
                'Kierunkowe efekty uczenia się (kody)',
                'Przedmiotowy efekt uczenia się',
                'Metody weryfikacji osiągnięcia efektu uczenia się'
            ]
            for i, h in enumerate(col_headers):
                _set_cell_bg(row.cells[i], COLOR_HEADER_BG)
                _run_in_cell(row.cells[i], h, bold=True, size=Pt(9))

            if items:
                for item in items:
                    row = table.add_row()
                    if isinstance(item, dict):
                        keu = str(item.get('keu') or '')
                        peu = str(item.get('peu') or '')
                        metoda = str(item.get('metoda_weryfikacji') or '')
                    else:
                        # fallback: stary format — item jest stringiem
                        keu = ''
                        peu = str(item)
                        metoda = ''
                    _value_cell(row.cells[0], keu)
                    _value_cell(row.cells[1], peu)
                    _value_cell(row.cells[2], metoda)
            else:
                row = table.add_row()
                _value_cell(row.cells[0], '')
                _value_cell(row.cells[1], '')
                _value_cell(row.cells[2], '')

    # -------------------------------------------------------------------------
    # Sekcja 21 — Wymagania laboratorium
    # -------------------------------------------------------------------------
    def _section_21(self):
        s = self.s
        wym = s.get('wymagania_laboratorium', {}) or {}
        pc_params  = wym.get('pc_params', []) or []
        software   = wym.get('software', []) or []
        wyposazenie = wym.get('wyposazenie_dodatkowe', []) or []

        table = self.doc.add_table(rows=0, cols=3)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        # Nagłówek scalony
        row = table.add_row()
        header_cell = row.cells[0]
        for i in range(1, 3):
            header_cell = header_cell.merge(row.cells[i])
        _set_cell_bg(header_cell, COLOR_HEADER_BG)
        _run_in_cell(header_cell, '21.  Wymagania dotyczące laboratorium, w którym odbywają się zajęcia z niniejszego przedmiotu', bold=True, size=FONT_SIZE_LABEL)

        # Nagłówki kolumn
        row = table.add_row()
        col_headers = [
            'Parametry dla stanowisk komputerowych',
            'Oprogramowanie konieczne do realizacji przedmiotu',
            'Wyposażenie specjalistyczne konieczne do realizacji przedmiotu'
        ]
        for i, h in enumerate(col_headers):
            _set_cell_bg(row.cells[i], COLOR_HEADER_BG)
            _run_in_cell(row.cells[i], h, bold=True, size=Pt(8))

        # Wiersze danych – każda pozycja listy w osobnym wierszu tabeli
        cols_data = [pc_params, software, wyposazenie]
        max_rows = max((len(c) for c in cols_data), default=1)
        if max_rows == 0:
            max_rows = 1
        for i in range(max_rows):
            row = table.add_row()
            for j, items in enumerate(cols_data):
                val = items[i] if i < len(items) else ''
                _value_cell(row.cells[j], val)

    # -------------------------------------------------------------------------
    # Sekcja 22 — Informacje dodatkowe
    # -------------------------------------------------------------------------
    def _section_22(self):
        s = self.s
        info = _safe(s, 'informacje_dodatkowe') or ''

        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '22.  Informacje dodatkowe – certyfikaty do których przygotowuje przedmiot', bold=True, size=FONT_SIZE_LABEL)

        row = table.add_row()
        _value_cell(row.cells[0], info)

    # -------------------------------------------------------------------------
    # Sekcja 23 — Uzasadnienie / Rynek pracy
    # -------------------------------------------------------------------------
    def _section_23(self):
        s = self.s
        rynek = s.get('rynek_pracy', {}) or {}
        dziedzina    = str(rynek.get('dziedzina_gospodarki') or '')
        zawody       = str(rynek.get('zawody') or '')
        prace_dypl   = rynek.get('prace_dyplomowe') or []

        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '23.  Uzasadnienie dla prowadzenia przedmiotu - współpraca z rynkiem pracy', bold=True, size=FONT_SIZE_LABEL)

        # A — dziedzina gospodarki
        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], 'A.  W jakiego typu firmach bądź dziedzinach gospodarki będą potrzebne umiejętności nabyte w trakcie zajęć:', bold=False, size=Pt(8))
        row = table.add_row()
        _value_cell(row.cells[0], dziedzina)

        # B — zawody
        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], 'B.  W jakich zawodach wiedza i umiejętności nabyte podczas zajęć są istotne:', bold=False, size=Pt(8))
        row = table.add_row()
        _value_cell(row.cells[0], zawody)

        # C — prace dyplomowe
        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], 'C.  Przykładowe tematy prac dyplomowych i projektów badawczych:', bold=False, size=Pt(8))
        row = table.add_row()
        cell = row.cells[0]
        if prace_dypl:
            first = True
            for i, pd in enumerate(prace_dypl):
                if first and cell.paragraphs and cell.paragraphs[0].text == '':
                    p = cell.paragraphs[0]
                    first = False
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(f'{i+1}. {pd}')
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL
        else:
            _value_cell(cell, '')

