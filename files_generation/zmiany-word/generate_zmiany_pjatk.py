# -*- coding: utf-8 -*-
"""
generate_zmiany_pjatk.py
========================
Generuje ostylowany dokument Word "Zestawienie zmian w programie studiow 2026/2027"
zgodnie z identyfikacja wizualna PJATK Filia w Gdansku (czerwien #B40000, szarosci).

Wymagania:
    pip install python-docx Pillow

Uzycie:
    python generate_zmiany_pjatk.py
    # wyjscie: program_studiow_2026-27_zmiany_PJATK.docx
"""

import io
import os
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# ---------------------------------------------------------------------------
# Kolory PJATK (z pjatk.sty)
# ---------------------------------------------------------------------------
PJATK_RED   = RGBColor(0xB4, 0x00, 0x00)   # #B40000
PJATK_GRAY  = RGBColor(0x50, 0x50, 0x50)   # ciemna szarość
PJATK_LGRAY = RGBColor(0xF0, 0xF0, 0xF0)   # jasna szarość (tło nagłówków tabel)
PJATK_DGRAY = RGBColor(0x32, 0x32, 0x32)   # bardzo ciemna szarość
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

# ---------------------------------------------------------------------------
# Ścieżki
# ---------------------------------------------------------------------------
BASE_DIR  = Path(__file__).parent.parent          # katalog pj-studies
LOGO_PDF  = Path(__file__).parent / "pjlogo.pdf"  # logo PJATK w PDF (obok skryptu)
LOGO_SVG  = BASE_DIR / "latex" / "pjatk-logo.svg"
OUT_FILE  = Path(__file__).parent / "program_studiow_2026-27_zmiany_PJATK.docx"

FONT_NAME = "Calibri"


# ===========================================================================
# Helpers – formatowanie XML / OOXML
# ===========================================================================

def _set_cell_bg(cell, rgb: RGBColor):
    """Ustawia kolor tła komórki tabeli (shading)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **edges):
    """
    Ustawia obramowanie komórki.
    edges = dict z kluczami: top, bottom, left, right, insideH, insideV
    wartość: dict z kluczami 'sz' (szerokość 1/8 pt), 'val', 'color'
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in edges.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"),   attrs.get("val",   "single"))
        el.set(qn("w:sz"),    str(attrs.get("sz", 4)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), attrs.get("color", "B40000"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _paragraph_border_bottom(paragraph, color="B40000", sz=6):
    """Dodaje poziomą linię pod paragrafem (używane dla nagłówków sekcji)."""
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_run(paragraph, text, bold=False, italic=False,
             size_pt=11, color: RGBColor = None, font=FONT_NAME):
    run = paragraph.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _set_spacing(paragraph, before_pt=0, after_pt=6, line_rule="auto", line=240):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after  = Pt(after_pt)


def _add_page_margins(doc, top=2.5, bottom=2.5, left=2.5, right=2.5):
    """Ustawia marginesy strony dla pierwszej (i jedynej) sekcji."""
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


# ===========================================================================
# Nagłówek / stopka
# ===========================================================================

def _setup_header_footer(doc):
    section = doc.sections[0]

    # --- nagłówek ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(hp, "Polsko-Japońska Akademia Technik Komputerowych  ·  Filia w Gdańsku",
             size_pt=9, color=PJATK_GRAY)
    _paragraph_border_bottom(hp, color="B40000", sz=4)
    hp.paragraph_format.space_after = Pt(4)

    # --- stopka ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # pole numeru strony
    run = fp.add_run()
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    run.font.color.rgb = PJATK_GRAY
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    _paragraph_border_bottom(fp, color="B40000", sz=4)  # linia nad stopką (wizualnie)


# ===========================================================================
# Konwersja logo → PNG
# ===========================================================================

def _logo_to_png_bytes(target_width_cm=5.0) -> bytes | None:
    """
    Konwertuje logo PJATK na PNG w kolejności:
      1. pjlogo.pdf  przez PyMuPDF  (plik obok skryptu)
      2. pjatk-logo.svg przez cairosvg
      3. pjatk-logo.svg przez Inkscape CLI
      4. Pillow – zastępcze logo tekstowe PJATK
    target_width_cm – docelowa szerokość w dokumencie (do skalowania DPI).
    Zwraca bytes PNG lub None.
    """
    # przelicz DPI: chcemy ~150 DPI końcowe (wystarczy dla logo), niezależnie od szerokości
    dpi = int(150 * target_width_cm / 5.0)

    # --- próba 1: pjlogo.pdf przez PyMuPDF ---
    if LOGO_PDF.exists():
        try:
            import fitz  # PyMuPDF
            doc  = fitz.open(str(LOGO_PDF))
            page = doc[0]
            # skaluj do ~1800 px szerokości (wystarczy na 5 cm @ 300 DPI)
            zoom = dpi / 72.0
            mat  = fitz.Matrix(zoom, zoom)
            pix  = page.get_pixmap(matrix=mat, alpha=False)
            png  = pix.tobytes("png")
            doc.close()
            print("[logo] pjlogo.pdf -> PNG przez PyMuPDF OK "
                  f"({pix.width}x{pix.height}px)")
            return png
        except Exception as e:
            print(f"[warn] PyMuPDF nieudany: {e}")

    # --- próba 2: cairosvg (wymaga libcairo-2.dll) ---
    if LOGO_SVG.exists():
        try:
            import cairosvg
            width_px = int(target_width_cm / 2.54 * 300)
            png = cairosvg.svg2png(url=str(LOGO_SVG), output_width=width_px)
            print("[logo] SVG -> PNG przez cairosvg OK")
            return png
        except Exception:
            pass

    # --- próba 3: Inkscape CLI ---
    if LOGO_SVG.exists():
        import subprocess, tempfile, shutil
        inkscape = shutil.which("inkscape")
        if inkscape:
            try:
                width_px = int(target_width_cm / 2.54 * 300)
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    tmp_path = tmp.name
                subprocess.run(
                    [inkscape, "--export-type=png",
                     f"--export-width={width_px}",
                     f"--export-filename={tmp_path}",
                     str(LOGO_SVG)],
                    check=True, capture_output=True,
                )
                data = Path(tmp_path).read_bytes()
                os.unlink(tmp_path)
                print("[logo] SVG -> PNG przez Inkscape OK")
                return data
            except Exception:
                pass

    # --- próba 4: Pillow – zastępcze logo tekstowe ---
    try:
        from PIL import Image, ImageDraw, ImageFont
        width_px = int(target_width_cm / 2.54 * 300)
        h = int(width_px * 0.3)
        img  = Image.new("RGBA", (width_px, h), (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        bar_h = int(h * 0.6)
        draw.rectangle([(0, 0), (width_px, bar_h)], fill=(180, 0, 0))
        try:
            fnt_big = ImageFont.truetype("calibrib.ttf", int(bar_h * 0.65))
        except Exception:
            fnt_big = ImageFont.load_default()
        draw.text((int(width_px * 0.04), int(bar_h * 0.15)), "PJATK",
                  font=fnt_big, fill=(255, 255, 255))
        try:
            fnt_sm = ImageFont.truetype("calibri.ttf", int((h - bar_h) * 0.5))
        except Exception:
            fnt_sm = ImageFont.load_default()
        draw.text((int(width_px * 0.02), bar_h + 3),
                  "Polsko-Japonska Akademia Technik Komputerowych  Filia w Gdansku",
                  font=fnt_sm, fill=(50, 50, 50))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        print("[logo] zastępcze logo Pillow OK")
        return buf.getvalue()
    except Exception as e:
        print(f"[warn] Nie mozna wygenerowac logo: {e}")
        return None


# ===========================================================================
# Strona tytułowa
# ===========================================================================

def _add_title_page(doc):
    # Logo PJATK – PDF preferowany, fallback SVG/Pillow
    # Szerokość = A4 (21 cm) - 2 × margines (2.5 cm) = 16 cm
    LOGO_WIDTH_CM = 16.0
    logo_png = None
    png_bytes = _logo_to_png_bytes(target_width_cm=LOGO_WIDTH_CM)
    if png_bytes:
        logo_png = io.BytesIO(png_bytes)

    if logo_png:
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_before = Pt(0)
        logo_p.paragraph_format.space_after  = Pt(16)
        run = logo_p.add_run()
        run.add_picture(logo_png, width=Cm(LOGO_WIDTH_CM))
    else:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(20)
        spacer.paragraph_format.space_after  = Pt(0)

    # Tytuł instytucji
    inst_p = doc.add_paragraph()
    inst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_p.paragraph_format.space_after = Pt(4)
    _add_run(inst_p, "Polsko-Japońska Akademia Technik Komputerowych",
             bold=True, size_pt=14, color=PJATK_DGRAY)

    inst_p2 = doc.add_paragraph()
    inst_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_p2.paragraph_format.space_after = Pt(30)
    _add_run(inst_p2, "Filia w Gdańsku",
             bold=False, size_pt=12, color=PJATK_GRAY)

    # Główny tytuł dokumentu
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after  = Pt(8)
    _add_run(title_p, "ZESTAWIENIE ZMIAN W PROGRAMIE STUDIÓW",
             bold=True, size_pt=20, color=PJATK_RED)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(30)
    _add_run(subtitle_p, "Rok akademicki 2026/2027",
             bold=False, size_pt=13, color=PJATK_GRAY)

    # Ramka z metadanymi (tabela 2-kolumnowa)
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.style = "Table Grid"

    meta_data = [
        ("Tryb studiów",       "Studia stacjonarne oraz niestacjonarne"),
        ("Zakres opracowania",
         "Zmiany godzinowe, przesunięcia przedmiotów między semestrami, "
         "nowe i usunięte przedmioty, zmiany w specjalizacjach"),
        ("Program studiów online", "https://pjmpr.github.io/GD_WI_PRG_26-27/"),
    ]

    for i, (label, value) in enumerate(meta_data):
        row = meta_table.rows[i]
        lc = row.cells[0]
        _set_cell_bg(lc, PJATK_LGRAY)
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(4)
        lp.paragraph_format.left_indent  = Cm(0.2)
        _add_run(lp, label, bold=True, size_pt=10, color=PJATK_DGRAY)
        vc = row.cells[1]
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(4)
        vp.paragraph_format.space_after  = Pt(4)
        vp.paragraph_format.left_indent  = Cm(0.2)
        # URL w ostatnim wierszu – czerwony, pogrubiony
        if i == len(meta_data) - 1:
            _add_run(vp, value, bold=True, size_pt=11, color=PJATK_RED)
        else:
            _add_run(vp, value, size_pt=10, color=PJATK_DGRAY)

    # obramowanie tabeli meta
    for row in meta_table.rows:
        for cell in row.cells:
            _set_cell_border(cell,
                top={"sz": 4, "color": "B40000"},
                bottom={"sz": 4, "color": "B40000"},
                left={"sz": 4, "color": "B40000"},
                right={"sz": 4, "color": "B40000"},
            )

    # szerokości kolumn: lewa węższa (4 cm), prawa szersza (12 cm)
    for row in meta_table.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(12.0)

    # rok / data
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(16)
    _add_run(date_p, "Gdańsk, 2026", size_pt=10, color=PJATK_GRAY)


    # podział strony
    doc.add_page_break()


# ===========================================================================
# Helpers – nagłówki sekcji
# ===========================================================================

def _add_section_heading(doc, text, level=1):
    """
    Dodaje nagłówek sekcji w stylu PJATK:
      level=1 → duży, czerwony, z linią
      level=2 → średni, ciemnoszary, bez linii
    """
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)
        _add_run(p, text, bold=True, size_pt=14, color=PJATK_RED)
        _paragraph_border_bottom(p, color="B40000", sz=6)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        _add_run(p, text, bold=True, size_pt=12, color=PJATK_DGRAY)
    return p


def _add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    _add_run(p, text, size_pt=11, color=PJATK_DGRAY)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before    = Pt(0)
    p.paragraph_format.space_after     = Pt(2)
    p.paragraph_format.left_indent     = Cm(0.5)
    _add_run(p, text, size_pt=10.5, color=PJATK_DGRAY)
    return p


# ===========================================================================
# Tabele
# ===========================================================================

def _add_styled_table(doc, headers, rows_data, col_widths_cm=None):
    """
    Tworzy tabelę z nagłówkiem (czerwone tło, biały tekst) i naprzemiennymi
    wierszami danych.
    headers: list[str]
    rows_data: list[list[str]]
    col_widths_cm: opcjonalna lista szerokości kolumn w cm
    """
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows_data), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # --- wiersz nagłówkowy ---
    hdr_row = table.rows[0]
    for j, hdr_text in enumerate(headers):
        cell = hdr_row.cells[j]
        _set_cell_bg(cell, PJATK_RED)
        _set_cell_border(cell,
            top={"sz": 4, "color": "B40000"},
            bottom={"sz": 4, "color": "B40000"},
            left={"sz": 4, "color": "B40000"},
            right={"sz": 4, "color": "B40000"},
        )
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hp = cell.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.space_before = Pt(3)
        hp.paragraph_format.space_after  = Pt(3)
        _add_run(hp, hdr_text, bold=True, size_pt=10, color=WHITE)

    # --- wiersze danych ---
    for i, row_data in enumerate(rows_data):
        bg = PJATK_LGRAY if i % 2 == 0 else WHITE
        tr = table.rows[i + 1]
        for j, cell_text in enumerate(row_data):
            cell = tr.cells[j]
            if i % 2 == 0:
                _set_cell_bg(cell, PJATK_LGRAY)
            _set_cell_border(cell,
                top={"sz": 2, "color": "C8C8C8", "val": "single"},
                bottom={"sz": 2, "color": "C8C8C8", "val": "single"},
                left={"sz": 2, "color": "C8C8C8", "val": "single"},
                right={"sz": 2, "color": "C8C8C8", "val": "single"},
            )
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after  = Pt(2)
            cp.paragraph_format.left_indent  = Cm(0.15)
            _add_run(cp, cell_text, size_pt=10, color=PJATK_DGRAY)

    # szerokości kolumn
    if col_widths_cm:
        for row in table.rows:
            for j, w in enumerate(col_widths_cm):
                row.cells[j].width = Cm(w)

    # mały odstęp po tabeli
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)

    return table


# ===========================================================================
# Dane dokumentu
# ===========================================================================

PRZENIESIENIA_HEADERS = ["Przedmiot", "Było\n(semestr)", "Jest\n(semestr)", "Uzasadnienie zmiany"]
PRZENIESIENIA_ROWS = [
    ["Analiza matematyczna",                     "2", "1", "Zmiana wynikająca z zaleceń PKA."],
    ["Algebra liniowa z geometrią",              "1", "2", "Zmiana wynikająca z zaleceń PKA."],
    ["Matematyka dyskretna",                     "1", "2", "Zmiana wynikająca z zaleceń PKA."],
    ["Historia i kultura Japonii",               "3", "1",
     "Poprawa zależności między przedmiotami, zwolnienie miejsca na inny przedmiot "
     "w semestrze 3 oraz lepsza reorganizacja programu studiów."],
    ["Technologie Internetu",                    "4", "2",
     "Przeniesienie na wcześniejszy semestr, aby student wcześniej zdobył kompetencje "
     "w zakresie tworzenia aplikacji internetowych i mógł szybciej budować portfolio projektów."],
    ["Fizyka",                                   "2", "3", "Lepsza reorganizacja przedmiotów w programie studiów."],
    ["Statystyczna analiza danych",              "7", "3",
     "Zmiana wynikająca z zaleceń PKA. Przedmiot kluczowy dla sztucznej inteligencji "
     "nie powinien być realizowany dopiero na ostatnim semestrze."],
    ["Systemy operacyjne",                       "4", "3", "Przeniesienie w celu lepszej reorganizacji programu studiów."],
    ["Bezpieczeństwo systemów informacyjnych",   "5", "4",
     "Przeniesienie na wcześniejszy semestr, aby przedmiot bazowy dla specjalizacji "
     "Cyberbezpieczeństwo był realizowany przed wyborem specjalizacji."],
    ["Interakcja człowiek–komputer",             "6", "5",
     "Przeniesienie przed rozpoczęciem projektów dyplomowych, aby studenci wcześniej "
     "poznali zasady projektowania interfejsów użytkownika."],
    ["Psychologia umiejętności inżynierskich",   "1", "5", "Lepsza reorganizacja przedmiotów w programie studiów."],
]

NOWE_USUNIETE_HEADERS = ["Semestr", "Zakres zmiany", "Przedmiot", "Uzasadnienie"]
NOWE_USUNIETE_ROWS = [
    ["1", "Nowe",
     "Warsztat programisty",
     "Odpowiedź na braki kompetencyjne w zakresie narzędzi programistycznych, debugowania, "
     "pracy zespołowej z Git, wykorzystania AI w pracy programisty oraz budowania portfolio "
     "projektów informatycznych."],
    ["1", "Usunięte",
     "Wstęp do architektury komputerów 1",
     "Przedmiot miał wyłącznie charakter wykładowy i był bardziej adekwatny dla innego profilu "
     "studiów; stanowił pozostałość po wcześniejszym profilu ogólnoakademickim."],
    ["2", "Nowe",
     "Języki programowania 1 i 2",
     "Wprowadzenie przedmiotów obieralnych umożliwiających poznanie różnych paradygmatów "
     "programowania oraz ich zastosowań w sztucznej inteligencji, grach, IoT, aplikacjach "
     "internetowych i analizie danych."],
    ["2", "Usunięte",
     "Warsztaty programistyczne",
     "Przedmiot koncentrował się wokół technologii PHP. Zdecydowano o zastąpieniu dotychczasowych "
     "rozproszonych treści jednym przedmiotem warsztatowym realizowanym przez cały semestr 3, "
     "obejmującym bazy danych, API oraz interfejs użytkownika."],
    ["3", "Nowe",
     "Programowanie aplikacji internetowych",
     "Przedmiot powstał z połączenia przedmiotów Metody programowania oraz Java zaawansowana, "
     "które powielały treści. Dodatkowo zmieniono technologię wiodącą z Java na C#."],
    ["3", "Nowe",
     "Systemy baz danych 1 i 2",
     "Wprowadzenie przedmiotów obieralnych umożliwiających poznanie różnych rodzajów baz danych "
     "NoSQL oraz ich praktycznych zastosowań."],
    ["3", "Usunięte",
     "Metody programowania oraz Java zaawansowana",
     "Przedmioty powielały istotną część treści programowych; utworzono nowy przedmiot "
     "z większą liczbą godzin przeznaczonych na wykład i laboratoria."],
    ["3", "Usunięte",
     "Systemy baz danych",
     "Zmiana wprowadzona w celu udostępnienia miejsca w programie na przedmioty obieralne "
     "z obszaru baz danych NoSQL."],
    ["4", "Nowe",
     "Programowanie systemowe",
     "Nowy przedmiot w miejsce Wstępu do architektury komputerów 2. Lepsze odzwierciedlenie "
     "treści programowych oraz wprowadzenie do programowania systemowego i IoT."],
    ["4", "Nowe",
     "Technologie aplikacji mobilnych 1 i 2",
     "Wprowadzenie przedmiotów obieralnych poświęconych tworzeniu aplikacji mobilnych, "
     "które wcześniej były marginalizowane w programie studiów."],
    ["5", "Nowe",
     "Kompetencje lidera IT",
     "Nowy przedmiot rozwijający kompetencje miękkie potrzebne w pracy lidera IT."],
    ["6", "Nowe",
     "Metody Design Thinking w projektowaniu systemów IT",
     "Wprowadzenie przedmiotu obieralnego rozwijającego kompetencje projektowe niezbędne "
     "w tworzeniu nowoczesnych systemów IT."],
]


# ===========================================================================
# Główna funkcja generująca dokument
# ===========================================================================

def build_document():
    doc = Document()

    # Marginesy
    _add_page_margins(doc, top=2.5, bottom=2.5, left=2.5, right=2.5)

    # Domyślna czcionka dokumentu
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.font.color.rgb = PJATK_DGRAY

    # Nagłówek / stopka
    _setup_header_footer(doc)

    # -----------------------------------------------------------------------
    # Strona tytułowa
    # -----------------------------------------------------------------------
    _add_title_page(doc)

    # -----------------------------------------------------------------------
    # 1. Zmiany w programie studiów stacjonarnych
    # -----------------------------------------------------------------------
    _add_section_heading(doc, "1.  Zmiany w programie studiów stacjonarnych", level=1)
    _add_styled_table(
        doc,
        headers=["", "Dotychczas", "Obecnie", "Uwagi"],
        rows_data=[
            ["Godziny dydaktyczne", "2715 godz.", "2710 godz.", "Zmiana o \u22125 godz."],
            ["Godziny BHP",         "0 godz.",    "4 godz.",    "Przeniesienie z puli dydaktycznej"],
            ["Suma \u0142\u0105czna",           "2715 godz.", "2714 godz.", "Zmiana o \u22121 godz."],
            ["Efekty uczenia si\u0119", "\u2014", "\u2014",
             "Efekty uczenia si\u0119 nie uleg\u0142y zmianie"],
        ],
        col_widths_cm=[4.5, 3.0, 3.0, 6.0],
    )

    # 1.1
    _add_section_heading(doc, "1.1.  Przeniesienia przedmiotów pomiędzy semestrami", level=2)
    _add_styled_table(
        doc,
        PRZENIESIENIA_HEADERS,
        PRZENIESIENIA_ROWS,
        col_widths_cm=[5.5, 2.2, 2.2, 6.6],
    )

    # 1.2
    doc.add_page_break()
    _add_section_heading(doc, "1.2.  Nowe oraz usunięte przedmioty", level=2)
    _add_styled_table(
        doc,
        NOWE_USUNIETE_HEADERS,
        NOWE_USUNIETE_ROWS,
        col_widths_cm=[1.8, 2.4, 4.8, 7.5],
    )

    # 1.3
    _add_section_heading(doc, "1.3.  Zmiany w specjalizacjach", level=2)

    # Specjalizacje
    specializations = [
        (
            "Aplikacje Internetowe \u2192 Architektury oprogramowania i DevOps \u2013 zmiana nazwy specjalizacji",
            [
                "Dotychczasowa specjalizacja Aplikacje Internetowe zosta\u0142a przemianowana na "
                "Architektury oprogramowania i DevOps.",
                "Semestr 5: Architektura mikroserwisowa i mikrofrontendowa; Konteneryzacja serwis\u00f3w internetowych.",
                "Semestr 6: Technologie DevOps.",
                "Semestr 7: Zarz\u0105dzanie infrastruktur\u0105 chmurowa.",
                "Pow\u00f3d: lepsza odpowied\u017a na rozw\u00f3j AI, potrzeby rynku pracy "
                "i zainteresowania student\u00f3w \u2013 przy zachowaniu ci\u0105g\u0142o\u015bci specjalizacji.",
            ],
        ),
        (
            "Inżynieria gier komputerowych – modyfikacja przedmiotów specjalizacyjnych",
            [
                'Semestr 5: zmiana nazwy przedmiotu \u201eWytwarzanie Gier 1\u201d na \u201ePrototypowanie gier komputerowych\u201d.',
                'Semestr 6: wprowadzenie nowego przedmiotu \u201eSilniki gier komputerowych\u201d.',
                'Semestr 7: zmiana nazwy przedmiotu \u201eWytwarzanie Gier 2\u201d na \u201eProjektowanie gier komputerowych\u201d.',
                "Powód: lepsze dopasowanie nazw do treści programowych oraz uzupełnienie specjalizacji "
                "o przedmiot poświęcony silnikom gier.",
            ],
        ),
        (
            "Sztuczna Inteligencja – zmiana nazwy przedmiotu",
            [
                'Semestr 5: zmiana nazwy przedmiotu \u201eFuture of deep learning\u201d na \u201eNowoczesne metody uczenia g\u0142\u0119bokiego\u201d.',
                "Powód: lepsze dopasowanie nazwy do treści programowych.",
            ],
        ),
    ]

    for spec_name, bullets in specializations:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(8)
        sp.paragraph_format.space_after  = Pt(2)
        _add_run(sp, spec_name, bold=True, size_pt=11, color=PJATK_RED)
        for bullet_text in bullets:
            _add_bullet(doc, bullet_text)

    # -----------------------------------------------------------------------
    # 2. Zmiany w programie studiów niestacjonarnych
    # -----------------------------------------------------------------------
    _add_section_heading(doc, "2.  Zmiany w programie studiów niestacjonarnych", level=1)
    _add_body_paragraph(doc, "Zmiany w godzinach: brak.")
    _add_body_paragraph(
        doc,
        "Charakter zmian: przyjęto rozwiązania analogiczne do programu studiów stacjonarnych.",
    )
    _add_body_paragraph(
        doc,
        "Uzasadnienie: ujednolicenie programów oraz wprowadzanych zmian dla studentów trybu "
        "niestacjonarnego z programem studiów stacjonarnych.",
    )

    # -----------------------------------------------------------------------
    # 3. Wniosek końcowy
    # -----------------------------------------------------------------------
    _add_section_heading(doc, "3.  Wniosek końcowy", level=1)
    _add_body_paragraph(
        doc,
        "Przedstawione zmiany mają charakter porządkujący, dostosowujący oraz rozwojowy. "
        "Obejmują one zarówno reorganizację kolejności przedmiotów (w celu poprawy logiki "
        "programu studiów oraz realizacji zaleceń PKA), jak i modernizację oferty dydaktycznej "
        "przez wprowadzenie nowych przedmiotów odpowiadających aktualnym potrzebom rynku pracy "
        "i rozwoju technologii.",
    )

    # -----------------------------------------------------------------------
    # Zapis
    # -----------------------------------------------------------------------
    doc.save(str(OUT_FILE))
    print(f"[OK] Dokument zapisany: {OUT_FILE}")


if __name__ == "__main__":
    build_document()

