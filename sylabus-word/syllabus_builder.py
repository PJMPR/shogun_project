# -*- coding: utf-8 -*-
"""
Klasa budująca dokument Word sylabusu na podstawie danych JSON.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import math
import json
from pathlib import Path

from styles import (
    COLOR_HEADER_BG, COLOR_HEADER_DARK, COLOR_WHITE, COLOR_BLACK,
    FONT_NAME, FONT_SIZE_NORMAL, FONT_SIZE_LABEL, FONT_SIZE_TITLE,
    FONT_SIZE_HEADER, MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT, MARGIN_RIGHT
)


def _safe(data, *keys, default=""):
    """Bezpieczne pobieranie zagnieżdżonych kluczy ze słownika."""
    current = data
    for key in keys:
        if isinstance(current, dict):
            current = current.get(key)
        else:
            return default
        if current is None:
            return default
    return current if current is not None else default


def _set_cell_bg(cell, color: str):
    """Ustawia kolor tła komórki tabeli. color = hex string np. 'D9D9D9'"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def _set_table_borders(table):
    """Ustawia obramowanie wszystkich komórek tabeli."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _set_table_width(table, doc):
    """Ustawia szerokość tabeli na 100% strony."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)


def _run_in_cell(cell, text, bold=False, size=None, italic=False):
    """Dodaje run do paragrafu komórki i formatuje go."""
    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size if size else FONT_SIZE_NORMAL
    run.font.bold = bold
    run.font.italic = italic
    return run


def _add_paragraph_to_cell(cell, text, bold=False, size=None, italic=False, first=False):
    """Dodaje nowy paragraf do komórki."""
    if first and cell.paragraphs and cell.paragraphs[0].text == '':
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size if size else FONT_SIZE_NORMAL
    run.font.bold = bold
    run.font.italic = italic
    return p


def _label_cell(cell, text):
    """Formatuje komórkę jako etykietę (szare tło, pogrubiona)."""
    _set_cell_bg(cell, COLOR_HEADER_BG)
    _run_in_cell(cell, text, bold=True, size=FONT_SIZE_LABEL)


def _value_cell(cell, text):
    """Formatuje komórkę jako wartość."""
    _run_in_cell(cell, str(text) if text else "", size=FONT_SIZE_NORMAL)


def _section_header_row(table, text, colspan=None):
    """Dodaje wiersz nagłówka sekcji (szare tło, pogrubiony, scalony)."""
    row = table.add_row()
    cell = row.cells[0]
    if colspan and len(row.cells) > 1:
        cell = row.cells[0]
        for i in range(1, len(row.cells)):
            cell.merge(row.cells[i])
    _set_cell_bg(cell, COLOR_HEADER_BG)
    _run_in_cell(cell, text, bold=True, size=FONT_SIZE_LABEL)
    return row


def _add_section_table(doc, header_text, num_cols=1):
    """Tworzy nową tabelę z nagłówkiem sekcji."""
    table = doc.add_table(rows=0, cols=num_cols)
    _set_table_borders(table)
    _set_table_width(table, doc)
    _section_header_row(table, header_text, colspan=num_cols)
    return table


def _merge_and_set(row, start_col, end_col):
    """Scala komórki od start_col do end_col (włącznie)."""
    cell = row.cells[start_col]
    for i in range(start_col + 1, end_col + 1):
        cell = cell.merge(row.cells[i])
    return cell


# ---------------------------------------------------------------------------
# Ładowanie efektów kierunkowych i budowa odwrotnego mapowania
# ---------------------------------------------------------------------------
_EFEKTY_FILE = Path(__file__).parent.parent / 'public' / 'assets' / 'efekty_ksztalcenia.json'
_efekty_cache: dict | None = None  # {kod_przedmiotu: {kategoria: [K_W01, ...]}}


def _load_efekty_kierunkowe() -> dict:
    """
    Wczytuje efekty_ksztalcenia.json i buduje odwrotne mapowanie:
        { "PAI": {"wiedza": ["K_W09", ...], "umiejetnosci": ["K_U05", ...], "kompetencje_spoleczne": ["K_K01", ...]} }
    """
    global _efekty_cache
    if _efekty_cache is not None:
        return _efekty_cache

    _efekty_cache = {}
    if not _EFEKTY_FILE.exists():
        return _efekty_cache

    try:
        with open(_EFEKTY_FILE, encoding='utf-8') as f:
            data = json.load(f)
    except Exception:
        return _efekty_cache

    efekty = data.get('efekty_ksztalcenia', {})

    # mapowanie nazw kategorii z pliku JSON na klucze używane w sylabus
    kategorie = {
        'wiedza': 'wiedza',
        'umiejetnosci': 'umiejetnosci',
        'kompetencje_spoleczne': 'kompetencje_spoleczne',
    }

    for kat_key, kat_label in kategorie.items():
        for efekt in efekty.get(kat_key, []):
            kod_efektu = efekt.get('kod_efektu', '')
            kody_przedmiotow = efekt.get('kody', [])
            for kod_p in kody_przedmiotow:
                kod_p_upper = str(kod_p).upper()
                if kod_p_upper not in _efekty_cache:
                    _efekty_cache[kod_p_upper] = {'wiedza': [], 'umiejetnosci': [], 'kompetencje_spoleczne': []}
                _efekty_cache[kod_p_upper][kat_label].append(kod_efektu)

    return _efekty_cache


# ---------------------------------------------------------------------------
# Mapowanie kod_przedmiotu → zakres specjalizacyjny
# ---------------------------------------------------------------------------
_ASSETS = Path(__file__).parent.parent / 'public' / 'assets'
_typ_cache: dict | None = None  # {kod_upper: 'obligatoryjny'|'obieralny'|'lektorat'|'specjalizacyjny'}


def _load_typ_przedmiotu() -> dict:
    """
    Buduje mapowanie kod_przedmiotu → zakres specjalizacyjny na podstawie:
      - program.json (obowiązkowe)
      - electives-other.json (obieralne, lektoraty)
      - electives-specializations.json (specjalizacyjne)
    Analogicznie dla katalogu niestacjonarne/.
    Priorytet: specjalizacyjny > lektorat > obieralny > obligatoryjny
    """
    global _typ_cache
    if _typ_cache is not None:
        return _typ_cache

    _typ_cache = {}

    def _process(assets_dir: Path):
        # Grupy lektoratowe — identyfikujemy po id grupy 'LEK*' lub label zawierającym 'lektorat'
        lektorat_ids: set = set()
        other_file = assets_dir / 'electives-other.json'
        if other_file.exists():
            try:
                with open(other_file, encoding='utf-8') as f:
                    other = json.load(f)
                for g in other.get('groups', []):
                    gid = str(g.get('id', '')).upper()
                    label = str(g.get('label', '')).lower()
                    if gid.startswith('LEK') or 'lektorat' in label:
                        lektorat_ids.add(gid)
                    for item in g.get('items', []):
                        kod = str(item.get('code', '')).upper()
                        if gid.startswith('LEK') or 'lektorat' in label:
                            _typ_cache[kod] = 'lektorat'
                        else:
                            # nie nadpisuj jeśli już ustawione jako lektorat
                            if _typ_cache.get(kod) != 'lektorat':
                                _typ_cache[kod] = 'obieralny'
            except Exception:
                pass

        # Specjalizacyjne
        spec_file = assets_dir / 'electives-specializations.json'
        if spec_file.exists():
            try:
                with open(spec_file, encoding='utf-8') as f:
                    spec = json.load(f)
                for sp in spec.get('specializations', []):
                    for item in sp.get('items', []):
                        kod = str(item.get('code', '')).upper()
                        _typ_cache[kod] = 'specjalizacyjny'  # nadpisuje obieralny
            except Exception:
                pass

        # Obowiązkowe z program.json — tylko jeśli kod nie jest jeszcze sklasyfikowany
        prog_file = assets_dir / 'program.json'
        if prog_file.exists():
            try:
                with open(prog_file, encoding='utf-8') as f:
                    prog = json.load(f)
                for sem in prog.get('semesters', []):
                    for s in sem.get('subjects', []):
                        kod = str(s.get('code', '')).upper()
                        if kod not in _typ_cache:
                            _typ_cache[kod] = 'obligatoryjny'
            except Exception:
                pass

    _process(_ASSETS)
    _process(_ASSETS / 'niestacjonarne')
    return _typ_cache


def _get_zakres_specjalizacyjny(kod: str) -> str:
    """Zwraca tekst zakresu specjalizacyjnego dla komórki w tabeli sekcji 7."""
    typ_map = _load_typ_przedmiotu()
    typ = typ_map.get(str(kod).upper(), 'obligatoryjny')
    mapping = {
        'specjalizacyjny': 'specjalizacyjny',
        'lektorat': 'lektorat',
        'obieralny': 'obieralny',
        'obligatoryjny': 'obligatoryjny',
    }
    return mapping.get(typ, 'obligatoryjny')


class SyllabusBuilder:
    """Buduje dokument Word sylabusu na podstawie danych z JSON."""

    def __init__(self, data: dict):
        self.data = data
        self.s = data.get('sylabus', {})
        self.doc = Document()
        self._setup_page()
        # Kody kierunkowych efektów dla tego przedmiotu
        kod = str(self.s.get('kod_przedmiotu', '')).upper()
        efekty_map = _load_efekty_kierunkowe()
        self._kody_kierunkowe = efekty_map.get(kod, {'wiedza': [], 'umiejetnosci': [], 'kompetencje_spoleczne': []})
        # Zakres specjalizacyjny i stopień
        self._zakres = _get_zakres_specjalizacyjny(kod)
        self._stopien = 'I'

    def _setup_page(self):
        """Konfiguruje marginesy strony."""
        for section in self.doc.sections:
            section.top_margin = MARGIN_TOP
            section.bottom_margin = MARGIN_BOTTOM
            section.left_margin = MARGIN_LEFT
            section.right_margin = MARGIN_RIGHT

    def build(self) -> Document:
        """Buduje kompletny dokument sylabusu."""
        self._add_title()
        self._section_1_6()
        self._section_7()
        self._section_8()
        self._section_9()
        self._section_10()
        self._section_11()
        self._section_12()
        self._section_13()
        self._section_14()
        self._section_15()
        self._section_16()
        self._section_17()
        self._section_18_20()
        self._section_21()
        self._section_22()
        self._section_23()
        return self.doc

    # -------------------------------------------------------------------------
    # Nagłówek dokumentu
    # -------------------------------------------------------------------------
    def _add_title(self):
        s = self.s
        uczelnia = _safe(s, 'uczelnia', default='Polsko-Japońska Akademia Technik Komputerowych')
        kierunek = _safe(s, 'kierunek', default='')

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run('SYLABUS')
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_TITLE
        run.font.bold = True

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        run2 = p2.add_run(uczelnia)
        run2.font.name = FONT_NAME
        run2.font.size = Pt(10)
        run2.font.bold = True

        if kierunek:
            p3 = self.doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after = Pt(4)
            run3 = p3.add_run(f'Kierunek: {kierunek}')
            run3.font.name = FONT_NAME
            run3.font.size = Pt(10)

    # -------------------------------------------------------------------------
    # Sekcja 1–6
    # -------------------------------------------------------------------------
    def _section_1_6(self):
        s = self.s
        table = self.doc.add_table(rows=0, cols=4)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        # Wiersz 1: etykiety
        row = table.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        _set_cell_bg(c0, COLOR_HEADER_BG)
        _run_in_cell(c0, '1.  Nazwa polska przedmiotu:', bold=True, size=FONT_SIZE_LABEL)

        c2 = row.cells[2].merge(row.cells[3])
        _set_cell_bg(c2, COLOR_HEADER_BG)
        _run_in_cell(c2, '2.  Wersja z dnia:', bold=True, size=FONT_SIZE_LABEL)

        # Wiersz 1: wartości
        row = table.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        _value_cell(c0, _safe(s, 'nazwa_przedmiotu'))

        c2 = row.cells[2].merge(row.cells[3])
        _value_cell(c2, _safe(s, 'wersja_z_dnia'))

        # Wiersz 2: etykiety
        row = table.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        _set_cell_bg(c0, COLOR_HEADER_BG)
        _run_in_cell(c0, '3.  Nazwa angielska przedmiotu:', bold=True, size=FONT_SIZE_LABEL)

        c2 = row.cells[2].merge(row.cells[3])
        _set_cell_bg(c2, COLOR_HEADER_BG)
        _run_in_cell(c2, '4.  Kod przedmiotu:', bold=True, size=FONT_SIZE_LABEL)

        # Wiersz 2: wartości
        row = table.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        _value_cell(c0, '')  # brak pola w JSON

        c2 = row.cells[2].merge(row.cells[3])
        _value_cell(c2, _safe(s, 'kod_przedmiotu'))

        # Wiersz 3: etykiety
        row = table.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        _set_cell_bg(c0, COLOR_HEADER_BG)
        _run_in_cell(c0, '5.  Nazwa jednostki prowadzącej przedmiot:', bold=True, size=FONT_SIZE_LABEL)

        c2 = row.cells[2].merge(row.cells[3])
        _set_cell_bg(c2, COLOR_HEADER_BG)
        _run_in_cell(c2, '6.  Liczba punktów ECTS:', bold=True, size=FONT_SIZE_LABEL)

        # Wiersz 3: wartości
        row = table.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        _value_cell(c0, _safe(s, 'jednostka'))

        c2 = row.cells[2].merge(row.cells[3])
        _value_cell(c2, str(_safe(s, 'ects')))

    # -------------------------------------------------------------------------
    # Sekcja 7 — Studia
    # -------------------------------------------------------------------------
    def _section_7(self):
        s = self.s
        table = self.doc.add_table(rows=0, cols=6)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        # Nagłówek scalony
        row = table.add_row()
        header_cell = row.cells[0]
        for i in range(1, 6):
            header_cell = header_cell.merge(row.cells[i])
        _set_cell_bg(header_cell, COLOR_HEADER_BG)
        _run_in_cell(header_cell, '7.  Studia', bold=True, size=FONT_SIZE_LABEL)

        # Wiersz nagłówkowy kolumn
        row = table.add_row()
        headers = ['kierunek', 'stopień', 'forma studiów', 'zakres specjalizacyjny', 'profil', 'semestr nauczania']
        for i, h in enumerate(headers):
            _set_cell_bg(row.cells[i], COLOR_HEADER_BG)
            _run_in_cell(row.cells[i], h, bold=True, size=Pt(8))

        # Wiersz danych
        tryb = _safe(s, 'tryb_studiow', default='')
        semestr = _safe(s, 'semestr_studiow', default='')
        profil = _safe(s, 'profil', default='')
        kierunek = _safe(s, 'kierunek', default='')
        values = [kierunek, self._stopien, tryb, self._zakres, profil, str(semestr)]
        row = table.add_row()
        for i, v in enumerate(values):
            _value_cell(row.cells[i], v)

    # -------------------------------------------------------------------------
    # Sekcja 8 — Osoba odpowiedzialna
    # -------------------------------------------------------------------------
    def _section_8(self):
        s = self.s
        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '8.  Osoba odpowiedzialna za przedmiot:', bold=True, size=FONT_SIZE_LABEL)

        row = table.add_row()
        _value_cell(row.cells[0], _safe(s, 'odpowiedzialny_za_przedmiot'))

    # -------------------------------------------------------------------------
    # Sekcja 9 — Formy zajęć
    # -------------------------------------------------------------------------
    def _section_9(self):
        s = self.s
        forma = s.get('forma_i_liczba_godzin_zajec', {})
        godziny = s.get('godziny', {})
        ects_total = _safe(s, 'ects', default=0)

        # Nagłówek sekcji
        t_header = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(t_header)
        _set_table_width(t_header, self.doc)
        row = t_header.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '9.  Formy zajęć, sposób ich realizacji i przypisana im liczba godzin:', bold=True, size=FONT_SIZE_LABEL)

        # Tabela godzin — 4 kolumny
        t_godziny = self.doc.add_table(rows=0, cols=4)
        _set_table_borders(t_godziny)
        _set_table_width(t_godziny, self.doc)

        row = t_godziny.add_row()
        col_headers = ['wykład', 'ćwiczenia/laboratorium/pracownia', 'wykład internetowo', 'ćw./lab./prac. internetowo']
        for i, h in enumerate(col_headers):
            _set_cell_bg(row.cells[i], COLOR_HEADER_BG)
            _run_in_cell(row.cells[i], h, bold=True, size=Pt(8))

        wyklady = forma.get('wyklady') or 0
        lab = forma.get('laboratorium_projekt') or 0
        cwiczenia = forma.get('cwiczenia_lektorat_seminarium') or 0
        total_cwiczenia = lab + cwiczenia

        row = t_godziny.add_row()
        _value_cell(row.cells[0], str(wyklady) if wyklady else '—')
        _value_cell(row.cells[1], str(total_cwiczenia) if total_cwiczenia else '—')
        _value_cell(row.cells[2], '—')
        _value_cell(row.cells[3], '—')

        # Tabela bilansu ECTS — 3 kolumny
        t_bilans = self.doc.add_table(rows=0, cols=3)
        _set_table_borders(t_bilans)
        _set_table_width(t_bilans, self.doc)

        row = t_bilans.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], 'Rodzaj nakładu pracy:', bold=True, size=Pt(8))
        _set_cell_bg(row.cells[1], COLOR_HEADER_BG)
        _run_in_cell(row.cells[1], 'Liczba godzin', bold=True, size=Pt(8))
        _set_cell_bg(row.cells[2], COLOR_HEADER_BG)
        _run_in_cell(row.cells[2], 'ECTS', bold=True, size=Pt(8))

        z_udzialem = godziny.get('z_udzialem_prowadzacego_h') or 0
        praca_wlasna = godziny.get('praca_wlasna_studenta_h') or 0
        calkowita = godziny.get('calkowita_liczba_godzin_h') or 0

        ects_z_udzialem = round(z_udzialem / 25, 1) if z_udzialem else 0
        ects_praca = round(praca_wlasna / 25, 1) if praca_wlasna else 0

        rows_bilans = [
            ('Zajęcia z bezpośrednim udziałem nauczyciela', str(z_udzialem), str(ects_z_udzialem)),
            ('Samodzielna praca studenta', str(praca_wlasna), str(ects_praca)),
            ('Suma', str(calkowita), str(ects_total)),
        ]
        for label, h, e in rows_bilans:
            row = t_bilans.add_row()
            _value_cell(row.cells[0], label)
            _value_cell(row.cells[1], h)
            _value_cell(row.cells[2], e)

    # -------------------------------------------------------------------------
    # Sekcja 10 — Język wykładowy
    # -------------------------------------------------------------------------
    def _section_10(self):
        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '10.  Język wykładowy:', bold=True, size=FONT_SIZE_LABEL)

        row = table.add_row()
        _value_cell(row.cells[0], 'polski')

    # -------------------------------------------------------------------------
    # Sekcja 11 — Metody dydaktyczne
    # -------------------------------------------------------------------------
    def _section_11(self):
        s = self.s
        metody = s.get('metody_dydaktyczne', {})

        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '11.  Metody dydaktyczne:', bold=True, size=FONT_SIZE_LABEL)

        row = table.add_row()
        cell = row.cells[0]
        all_methods = []

        if isinstance(metody, dict):
            for key, methods in metody.items():
                if methods and isinstance(methods, list):
                    label = 'Wykład' if 'wyklad' in key.lower() or 'wykład' in key.lower() else key.capitalize()
                    all_methods.append(('label', f'{label}:'))
                    for m in methods:
                        all_methods.append(('item', f'• {m}'))
        elif isinstance(metody, list):
            for m in metody:
                all_methods.append(('item', f'• {m}'))

        first = True
        for kind, text in all_methods:
            if first and cell.paragraphs and cell.paragraphs[0].text == '':
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_NORMAL
            run.font.bold = (kind == 'label')

    # -------------------------------------------------------------------------
    # Sekcja 12 — Zaliczenie i kryteria oceny
    # -------------------------------------------------------------------------
    def _section_12(self):
        s = self.s
        zaliczenie = s.get('zaliczenie', {})
        kryteria = s.get('kryteria_oceny', [])

        table = self.doc.add_table(rows=0, cols=2)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        # Nagłówek scalony
        row = table.add_row()
        header_cell = row.cells[0].merge(row.cells[1])
        _set_cell_bg(header_cell, COLOR_HEADER_BG)
        _run_in_cell(header_cell, '12.  Forma i sposób zaliczenia oraz podstawowe kryteria oceny lub wymagania egzaminacyjne:', bold=True, size=FONT_SIZE_LABEL)

        # Nagłówki kolumn
        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], 'wykład', bold=True, size=Pt(8))
        _set_cell_bg(row.cells[1], COLOR_HEADER_BG)
        _run_in_cell(row.cells[1], 'ćwiczenia/laboratorium/pracownia', bold=True, size=Pt(8))

        # Formy zaliczenia
        forms = ['Nieoceniany', 'Zaliczenie bez oceny', 'Zaliczenie z oceną', 'Egzamin']
        wyklad_sposob = _safe(zaliczenie, 'Wykład', 'sposob')
        lab_sposob = _safe(zaliczenie, 'Laboratorium', 'sposob')

        row = table.add_row()
        for col_idx, sposob in enumerate([wyklad_sposob, lab_sposob]):
            cell = row.cells[col_idx]
            first = True
            for f in forms:
                if first and cell.paragraphs and cell.paragraphs[0].text == '':
                    p = cell.paragraphs[0]
                    first = False
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                mark = '☑' if f == sposob else '☐'
                run = p.add_run(f'{mark} {f}')
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL

        # Etykiety kryteriów
        row = table.add_row()
        for col_idx in range(2):
            _set_cell_bg(row.cells[col_idx], COLOR_HEADER_BG)
            _run_in_cell(row.cells[col_idx], 'kryteria oceny:', bold=True, size=Pt(8))

        # Wartości kryteriów
        row = table.add_row()
        for col_idx in range(2):
            cell = row.cells[col_idx]
            if kryteria:
                first = True
                for i, k in enumerate(kryteria):
                    if first and cell.paragraphs and cell.paragraphs[0].text == '':
                        p = cell.paragraphs[0]
                        first = False
                    else:
                        p = cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
                    run = p.add_run(f'{i+1}. {k}')
                    run.font.name = FONT_NAME
                    run.font.size = FONT_SIZE_NORMAL
            else:
                _value_cell(cell, '')

    # -------------------------------------------------------------------------
    # Sekcja 13 — Przedmioty wprowadzające
    # -------------------------------------------------------------------------
    def _section_13(self):
        s = self.s
        przedmioty = s.get('przedmioty_wprowadzajace', [])

        table = self.doc.add_table(rows=0, cols=2)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        # Nagłówek scalony
        row = table.add_row()
        header_cell = row.cells[0].merge(row.cells[1])
        _set_cell_bg(header_cell, COLOR_HEADER_BG)
        _run_in_cell(header_cell, '13.  Kompetencje niezbędne dla realizacji przedmiotu i przedmioty poprzedzające, z których one pochodzą:', bold=True, size=FONT_SIZE_LABEL)

        # Nagłówki kolumn
        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], 'Nazwa przedmiotu poprzedzającego lub realizowanego równolegle', bold=True, size=Pt(8))
        _set_cell_bg(row.cells[1], COLOR_HEADER_BG)
        _run_in_cell(row.cells[1], 'Wymagane kompetencje niezbędne w realizacji danego przedmiotu', bold=True, size=Pt(8))

        if przedmioty:
            for p_item in przedmioty:
                row = table.add_row()
                _value_cell(row.cells[0], _safe(p_item, 'nazwa'))
                _value_cell(row.cells[1], _safe(p_item, 'wymagania'))
        else:
            row = table.add_row()
            _value_cell(row.cells[0], '')
            _value_cell(row.cells[1], '')

    # -------------------------------------------------------------------------
    # Sekcja 14 — Cel dydaktyczny
    # -------------------------------------------------------------------------
    def _section_14(self):
        s = self.s
        cel_pl = _safe(s, 'cel_dydaktyczny')
        cel_eng = _safe(s, 'cel_dydaktyczny_eng')

        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '14.  Cele przedmiotu (abstrakt w języku polskim i angielskim)', bold=True, size=FONT_SIZE_LABEL)

        row = table.add_row()
        cell = row.cells[0]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run('Abstrakt po polsku:')
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_NORMAL
        run.font.bold = True
        if cel_pl:
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_before = Pt(1)
            p2.paragraph_format.space_after = Pt(4)
            run2 = p2.add_run(cel_pl)
            run2.font.name = FONT_NAME
            run2.font.size = FONT_SIZE_NORMAL

        p3 = cell.add_paragraph()
        p3.paragraph_format.space_before = Pt(4)
        p3.paragraph_format.space_after = Pt(1)
        run3 = p3.add_run('Abstrakt po angielsku:')
        run3.font.name = FONT_NAME
        run3.font.size = FONT_SIZE_NORMAL
        run3.font.bold = True
        if cel_eng:
            p4 = cell.add_paragraph()
            p4.paragraph_format.space_before = Pt(1)
            p4.paragraph_format.space_after = Pt(1)
            run4 = p4.add_run(cel_eng)
            run4.font.name = FONT_NAME
            run4.font.size = FONT_SIZE_NORMAL
            run4.font.italic = True

    # -------------------------------------------------------------------------
    # Sekcja 15 — Treści programowe
    # -------------------------------------------------------------------------
    def _section_15(self):
        s = self.s
        tresci = s.get('tresci_programowe', [])

        table = self.doc.add_table(rows=0, cols=3)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        # Nagłówek scalony
        row = table.add_row()
        header_cell = row.cells[0]
        for i in range(1, 3):
            header_cell = header_cell.merge(row.cells[i])
        _set_cell_bg(header_cell, COLOR_HEADER_BG)
        _run_in_cell(header_cell, '15.  Treści programowe poszczególnych zajęć:', bold=True, size=FONT_SIZE_LABEL)

        # Nagłówki kolumn
        row = table.add_row()
        col_headers = ['Nr zajęć', 'Wykład', 'Ćwiczenia / Laboratorium / Pracownia']
        for i, h in enumerate(col_headers):
            _set_cell_bg(row.cells[i], COLOR_HEADER_BG)
            _run_in_cell(row.cells[i], h, bold=True, size=Pt(8))

        if tresci:
            for i, tresc in enumerate(tresci):
                row = table.add_row()
                _value_cell(row.cells[0], f'{i+1}.')
                _value_cell(row.cells[1], tresc)
                _value_cell(row.cells[2], '')
        else:
            for i in range(1, 9):
                row = table.add_row()
                _value_cell(row.cells[0], f'{i}.')
                _value_cell(row.cells[1], '')
                _value_cell(row.cells[2], '')

    # -------------------------------------------------------------------------
    # Sekcja 16 — Elektroniczne materiały dydaktyczne
    # -------------------------------------------------------------------------
    def _section_16(self):
        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '16.  Elektroniczne materiały dydaktyczne', bold=True, size=FONT_SIZE_LABEL)

        row = table.add_row()
        _value_cell(row.cells[0], '')

    # -------------------------------------------------------------------------
    # Sekcja 17 — Literatura
    # -------------------------------------------------------------------------
    def _section_17(self):
        s = self.s
        literatura = s.get('literatura', {})
        podstawowa = _safe(literatura, 'podstawowa', 'pozycje', default=[])
        uzupelniajaca = _safe(literatura, 'uzupelniajaca', 'pozycje', default=[])
        dokumentacja = literatura.get('dokumentacja_internetowa', {})

        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '17.  Wykaz literatury', bold=True, size=FONT_SIZE_LABEL)

        # Literatura wymagana
        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], 'Literatura wymagana do ostatecznego zaliczenia zajęć, wykorzystywana podczas zajęć oraz studiowana samodzielnie przez studenta:', bold=False, size=Pt(8))

        row = table.add_row()
        cell = row.cells[0]
        if podstawowa:
            first = True
            for i, poz in enumerate(podstawowa):
                if first and cell.paragraphs and cell.paragraphs[0].text == '':
                    p = cell.paragraphs[0]
                    first = False
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(f'{i+1}. {poz}')
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL
        else:
            _value_cell(cell, '')

        # Literatura uzupełniająca
        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], 'Literatura uzupełniająca:', bold=False, size=Pt(8))

        row = table.add_row()
        cell = row.cells[0]
        if uzupelniajaca:
            first = True
            for i, poz in enumerate(uzupelniajaca):
                if first and cell.paragraphs and cell.paragraphs[0].text == '':
                    p = cell.paragraphs[0]
                    first = False
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(f'{i+1}. {poz}')
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL
        else:
            _value_cell(cell, '')

        # Dokumentacja internetowa
        if dokumentacja:
            row = table.add_row()
            _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
            _run_in_cell(row.cells[0], 'Dokumentacja internetowa:', bold=False, size=Pt(8))

            row = table.add_row()
            cell = row.cells[0]
            first = True
            for i, (k, v) in enumerate(dokumentacja.items()):
                tekst = v if v else k
                if first and cell.paragraphs and cell.paragraphs[0].text == '':
                    p = cell.paragraphs[0]
                    first = False
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(f'• {tekst}')
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL

    # -------------------------------------------------------------------------
    # Sekcje 18–20 — Efekty kształcenia
    # -------------------------------------------------------------------------
    def _section_18_20(self):
        s = self.s
        efekty = s.get('efekty_ksztalcenia', {})

        sections = [
            ('18.  Wiedza nabywana / dostarczana uczestnikom w trakcie realizacji przedmiotu', 'wiedza'),
            ('19.  Umiejętności nabywane podczas realizacji przedmiotu', 'umiejetnosci'),
            ('20.  Kompetencje społeczne (postawy) kształtowane podczas realizacji przedmiotu', 'kompetencje_spoleczne'),
        ]

        for section_title, key in sections:
            items = efekty.get(key, [])
            # Kody kierunkowych efektów dla tej kategorii
            kody_kierunkowe = self._kody_kierunkowe.get(key, [])

            table = self.doc.add_table(rows=0, cols=3)
            _set_table_borders(table)
            _set_table_width(table, self.doc)

            # Nagłówek scalony
            row = table.add_row()
            header_cell = row.cells[0]
            for i in range(1, 3):
                header_cell = header_cell.merge(row.cells[i])
            _set_cell_bg(header_cell, COLOR_HEADER_BG)
            _run_in_cell(header_cell, section_title, bold=True, size=FONT_SIZE_LABEL)

            # Nagłówki kolumn
            row = table.add_row()
            col_headers = [
                'Kierunkowe efekty uczenia się (kody)',
                'Przedmiotowy efekt uczenia się',
                'Metody weryfikacji osiągnięcia efektu uczenia się'
            ]
            for i, h in enumerate(col_headers):
                _set_cell_bg(row.cells[i], COLOR_HEADER_BG)
                _run_in_cell(row.cells[i], h, bold=True, size=Pt(9))

            if items:
                for idx, item in enumerate(items):
                    row = table.add_row()
                    # Kod kierunkowy — przypisz kolejny z listy (jeśli dostępny)
                    kod_kier = kody_kierunkowe[idx] if idx < len(kody_kierunkowe) else ''
                    _value_cell(row.cells[0], kod_kier)
                    _value_cell(row.cells[1], item)
                    _value_cell(row.cells[2], '')
            else:
                # Brak efektów przedmiotowych — pokaż tylko kody kierunkowe (jeśli są)
                if kody_kierunkowe:
                    for kod_kier in kody_kierunkowe:
                        row = table.add_row()
                        _value_cell(row.cells[0], kod_kier)
                        _value_cell(row.cells[1], '')
                        _value_cell(row.cells[2], '')
                else:
                    for _ in range(3):
                        row = table.add_row()
                        _value_cell(row.cells[0], '')
                        _value_cell(row.cells[1], '')
                        _value_cell(row.cells[2], '')

    # -------------------------------------------------------------------------
    # Sekcja 21 — Wymagania laboratorium
    # -------------------------------------------------------------------------
    def _section_21(self):
        table = self.doc.add_table(rows=0, cols=3)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        # Nagłówek scalony
        row = table.add_row()
        header_cell = row.cells[0]
        for i in range(1, 3):
            header_cell = header_cell.merge(row.cells[i])
        _set_cell_bg(header_cell, COLOR_HEADER_BG)
        _run_in_cell(header_cell, '21.  Wymagania dotyczące laboratorium, w którym odbywają się zajęcia z niniejszego przedmiotu', bold=True, size=FONT_SIZE_LABEL)

        # Nagłówki kolumn
        row = table.add_row()
        col_headers = [
            'Parametry dla stanowisk komputerowych',
            'Oprogramowanie konieczne do realizacji przedmiotu',
            'Wyposażenie specjalistyczne konieczne do realizacji przedmiotu'
        ]
        for i, h in enumerate(col_headers):
            _set_cell_bg(row.cells[i], COLOR_HEADER_BG)
            _run_in_cell(row.cells[i], h, bold=True, size=Pt(8))

        # Puste wiersze
        for _ in range(2):
            row = table.add_row()
            for j in range(3):
                _value_cell(row.cells[j], '')

    # -------------------------------------------------------------------------
    # Sekcja 22 — Informacje dodatkowe
    # -------------------------------------------------------------------------
    def _section_22(self):
        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '22.  Informacje dodatkowe – certyfikaty do których przygotowuje przedmiot', bold=True, size=FONT_SIZE_LABEL)

        row = table.add_row()
        _value_cell(row.cells[0], '')

    # -------------------------------------------------------------------------
    # Sekcja 23 — Uzasadnienie
    # -------------------------------------------------------------------------
    def _section_23(self):
        table = self.doc.add_table(rows=0, cols=1)
        _set_table_borders(table)
        _set_table_width(table, self.doc)

        row = table.add_row()
        _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
        _run_in_cell(row.cells[0], '23.  Uzasadnienie dla prowadzenia przedmiotu - współpraca z rynkiem pracy', bold=True, size=FONT_SIZE_LABEL)

        sub_sections = [
            'A.  W jakiego typu firmach bądź dziedzinach gospodarki będą potrzebne umiejętności nabyte w trakcie zajęć:',
            'B.  W jakich zawodach wiedza i umiejętności nabyte podczas zajęć są istotne:',
            'C.  Przykładowe tematy prac dyplomowych i projektów badawczych:'
        ]
        for sub in sub_sections:
            row = table.add_row()
            _set_cell_bg(row.cells[0], COLOR_HEADER_BG)
            _run_in_cell(row.cells[0], sub, bold=False, size=Pt(8))

            row = table.add_row()
            _value_cell(row.cells[0], '')

